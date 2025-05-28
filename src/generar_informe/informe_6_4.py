
import json
import os
import sys
import re
import pandas as pd
import requests
from docx import Document
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt  
from dotenv import load_dotenv
from datetime import datetime
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl import load_workbook
from config import cargar_configuracion



# Obtener la ruta absoluta del directorio `src`
SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# Agregar `src` al sys.path 
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)

# Variable global
base_dir = os.path.dirname(__file__)
FECHA_INICIO_COL = 'Fecha Inicio'
FECHA_FIN_COL = 'Fecha Fin'
DURATION_COL = 'Duration'
COLOR_TOTAL = "BDD7EE"
COLOR_SUSTAIN = "C6E0B4"
COLOR_TOTAL_DATA = "DCE6F1"
COLOR_SUSTAIN_DATA = "EBF1DE"
FONT_PINK = Font(color="8A2BE2", bold=True)
FONT_BOLD = Font(bold=True)
ALIGN_CENTER = Alignment(horizontal="center", vertical="center")
TOTAL_RESEARCH_LABEL = "Total research"
SUSTAIN_RESEARCH_LABEL = "Sustainability research"
USD_LABEL = "US Dollar Spot Exchange Rates"
CURRENCY_INDEX = ["Euros", USD_LABEL, "Dolars"]
RESUMEN_FINANCIERO="Resumen Financiero"




load_dotenv()
# Configuración de la API
config = cargar_configuracion()
base_url = config["base_url"]
api_key = config["api_key"]
myModel = config["model"]

def preparar_fechas(df):
    """
    Prepara las columnas de fecha de inicio y fin en el DataFrame.

    Convierte las columnas 'Fecha Inicio' y 'Fecha Fin' a formato datetime,
    calcula la duración en días y verifica si hay filas con fechas vacías.

    Args:
        df (pd.DataFrame): DataFrame de entrada con datos de proyectos.

    Returns:
        pd.DataFrame: DataFrame con columnas de fecha convertidas y duración calculada.
                      Retorna el DataFrame original si se detectan filas con fechas vacías.
    """
    # Conversión vectorizada (rápida y garantiza que las columnas sean datetime64)
    df[FECHA_INICIO_COL] = pd.to_datetime(df[FECHA_INICIO_COL], dayfirst=True, errors='coerce')
    df[FECHA_FIN_COL] = pd.to_datetime(df[FECHA_FIN_COL], dayfirst=True, errors='coerce')

    # Verificar si alguna fila tiene ambas fechas vacías
    filas_vacias = df[df[FECHA_INICIO_COL, FECHA_FIN_COL]].isna().all(axis=1)
    if not filas_vacias.empty:
        print("Se encontraron filas con ambas fechas vacías. Deteniendo el procesamiento.")
        return df

    # Calcular la duración de forma vectorizada
    df[DURATION_COL] = (df[FECHA_FIN_COL] - df[FECHA_INICIO_COL]).dt.days + 1

    return df

def calcular_imputacion_diaria(df):
    """
    Calcula la imputación diaria para cada proyecto.

    Convierte la columna 'CUANTÍA TOTAL' a float y divide por la duración.

    Args:
        df (pd.DataFrame): DataFrame con datos de proyectos, incluyendo 'CUANTÍA TOTAL' y 'Duration'.

    Returns:
        pd.DataFrame: DataFrame con la columna 'Daily Imputation' añadida.
    """
    df['CUANTÍA TOTAL'] = (
        df['CUANTÍA TOTAL']
        .astype(str)
        .str.replace('.', '', regex=False)
        .str.replace(',', '.', regex=False)
        .astype(float)
    )
    df['Daily Imputation'] = df['CUANTÍA TOTAL'] / df[DURATION_COL]
    return df

def imputar_por_anho(df):
    """
    Imputa la cuantía total de cada proyecto a los años que abarca.

    Calcula la proporción de días de cada proyecto que caen en cada año
    y asigna la imputación diaria a esos días. Crea columnas para cada año
    y para la imputación sostenible por año.

    Args:
        df (pd.DataFrame): DataFrame con datos de proyectos, incluyendo fechas,
                           duración e imputación diaria.

    Returns:
        pd.DataFrame: DataFrame con columnas añadidas para la imputación anual
                      y la imputación sostenible anual.
    """
    if df.empty:
        print("No hay datos válidos para procesar.")
        return df

    # Determinar rango de años
    years = range(df[FECHA_INICIO_COL].dt.year.min(), df[FECHA_INICIO_COL].dt.year.max() + 1)

    # Inicializar columnas
    for year in years:
        df[str(year)] = 0.0
        df[f"{year}_sostenible"] = 0.0

    # Calcular imputación anual y llenar las columnas
    for i, row in df.iterrows():
        if pd.isna(row[FECHA_INICIO_COL]) or pd.isna(row[FECHA_FIN_COL]):
            print(f"Fila {i} tiene fechas vacías. Saltando.")
            continue

        for year in years:
            year_start = datetime(year, 1, 1)
            year_end = datetime(year, 12, 31)
            effective_start = max(row[FECHA_INICIO_COL], year_start)
            effective_end = min(row[FECHA_FIN_COL], year_end)

            if effective_start <= effective_end:
                days_in_year = (effective_end - effective_start).days + 1
                imputacion = round(days_in_year * row['Daily Imputation'], 2)

                df.at[i, str(year)] = imputacion

                if str(row.get("Sostenible", "no")).lower().strip() == "yes":
                    df.at[i, f"{year}_sostenible"] = imputacion

    return df

def es_sostenible_api(excel):
    """
    Determina si un proyecto es sostenible utilizando una API externa.

    Envía el título del proyecto a una API de clasificación y devuelve 'yes' o 'no'.

    Args:
        excel (str): El título del proyecto a clasificar.

    Returns:
        str | None: 'yes' si el proyecto es sostenible, 'no' si no lo es o hay un error,
                    o None si el título de entrada está vacío.
    """
    if not excel:
        return None

    body = {
        "model": myModel,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a helpful assistant that classifies if a project is related to sustainability. Answer only 'yes' or 'no'\n"
                    "1. Fin de la pobreza\n"
                    "2. Hambre cero\n"
                    "3. Salud y bienestar\n"
                    "4. Educación de calidad\n"
                    "5. Igualdad de género\n"
                    "6. Agua limpia y saneamiento\n"
                    "7. Energía asequible y no contaminante\n"
                    "8. Trabajo decente y crecimiento económico\n"
                    "9. Industria, innovación e infraestructura\n"
                    "10. Reducción de las desigualdades\n"
                    "11. Ciudades y comunidades sostenibles\n"
                    "12. Producción y consumo responsables\n"
                    "13. Acción por el clima\n"
                    "14. Vida submarina\n"
                    "15. Vida de ecosistemas terrestres\n"
                    "16. Paz, justicia e instituciones sólidas\n"
                    "17. Alianzas para lograr los objetivos\n"
                )
            },
            {
                "role": "user",
                "content": excel
            }
        ],
        "temperature": 0.2
    }

    body_json = json.dumps(body)

    response = requests.post(
            f"{base_url}/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            data=body_json
        )

    if response.status_code != 200:
        print(f"Error en la API para {excel}: {response.status_code} - {response.text}")
        return "no"
       

    #  Obtener respuesta limpia
    message_content = response.json()['choices'][0]['message']['content'].strip().lower()
    # Filtrar solo el contenido posterior a </think>
    if "</think>" in message_content:
        message_content = message_content.split("</think>")[-1].strip()  # Solo lo que sigue después de </think>
    return message_content

def marcar_sostenibles_api(df, columna_titulo='Título'):
    """
    Marca los proyectos en el DataFrame como sostenibles o no utilizando la API.

    Aplica la función `es_sostenible_api` a la columna especificada (por defecto 'Título')
    y guarda el resultado en una nueva columna 'Sostenible'.

    Args:
        df (pd.DataFrame): DataFrame con datos de proyectos.
        columna_titulo (str, optional): Nombre de la columna que contiene el título del proyecto.
                                        Por defecto es 'Título'.

    Returns:
        pd.DataFrame: DataFrame con la columna 'Sostenible' añadida.
    """
    df['Sostenible'] = df[columna_titulo].apply(lambda x: es_sostenible_api(x))
    return df      
   
def cargar_excel_con_header_dinamico(path, columnas_esperadas):
    """
    Carga un archivo Excel buscando dinámicamente la fila de encabezado.

    Busca la primera fila que contiene todas las `columnas_esperadas`, la usa como
    encabezado, limpia los nombres de columna, elimina columnas sin nombre y
    elimina filas después de una fila que contenga "Total:".

    Args:
        path (str): Ruta al archivo Excel.
        columnas_esperadas (list): Lista de nombres de columnas que se esperan encontrar
                                   en la fila de encabezado.

    Returns:
        pd.DataFrame: DataFrame limpio con el encabezado correcto.

    Raises:
        ValueError: Si no se encuentran todas las columnas esperadas en ninguna fila.
    """

    df_bruto = pd.read_excel(path, header=None)

    for i in range(len(df_bruto)):
        posibles_columnas = df_bruto.iloc[i].tolist()

        if all(col in posibles_columnas for col in columnas_esperadas):
            df = pd.read_excel(path, header=i)

            # Limpiar nombres de columnas
            df.columns = df.columns.str.strip()

            # Eliminar columnas tipo "Unnamed"
            df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
            
            # Detectar fila con "Total:" antes de limpiar
            for j, row in df.iterrows():
                if any(str(cell).strip().lower() == "total:" for cell in row):
                    df = df.iloc[:j]
                    break

            # Eliminar NaNs después de cortar el DataFrame
            df_limpio = df.dropna(subset=columnas_esperadas)
            
            return df_limpio

    raise ValueError(" No se encontraron todas las columnas esperadas en ninguna fila.")

def anhadir_fila_total(df):
    """
    Añade una fila de totales al final del DataFrame.

    Suma los valores de las columnas numéricas (años y años sostenibles)
    y crea una nueva fila con estos totales, etiquetada como 'TOTAL '.

    Args:
        df (pd.DataFrame): DataFrame con datos de proyectos y columnas de imputación anual.

    Returns:
        pd.DataFrame: DataFrame con la fila de totales añadida al final.
    """
    
    # Columnas que contienen números (años y sostenibles)
    columnas_suma = [col for col in df.columns if col.isdigit() or col.endswith('_sostenible')]

    # Creamos la fila vacía
    fila_vacia = pd.Series({col: "" for col in df.columns})

    # Creamos la fila de suma
    fila_suma = {}
    for col in df.columns:
        if col in columnas_suma:
            fila_suma[col] = df[col].sum()
        else:
            fila_suma[col] = ""
            
    primera_columna= df.columns[0]
    fila_suma[primera_columna]= 'TOTAL '

    fila_suma = pd.Series(fila_suma)

    # Añadimos ambas filas al final del DataFrame
    df = pd.concat([df, pd.DataFrame([fila_vacia]), pd.DataFrame([fila_suma])], ignore_index=True)

    return df

def duplicar_valores_sostenibles(df):
    """
    Copia los valores de imputación anual a las columnas sostenibles
    para los proyectos marcados como 'yes' en la columna 'Sostenible'.

    Args:
        df (pd.DataFrame): DataFrame con columnas de imputación anual y 'Sostenible'.

    Returns:
        pd.DataFrame: DataFrame con los valores de imputación sostenible actualizados.
    """

    # Años detectados (solo columnas con nombres numéricos)
    columnas_anhos = [col for col in df.columns if col.isdigit()]
    
    # Columnas sostenibles correspondientes
    columnas_sostenibles = [f"{anho}_sostenible" for anho in columnas_anhos if f"{anho}_sostenible" in df.columns]

    for anho, anho_sostenible in zip(columnas_anhos, columnas_sostenibles):
        df.loc[df["Sostenible"].str.lower() == "yes", anho_sostenible] = df.loc[df["Sostenible"].str.lower() == "yes", anho]

    return df

def exportar_resultado(df, path_salida_temp):
    """
    Copia los valores de imputación anual a las columnas sostenibles
    para los proyectos marcados como 'yes' en la columna 'Sostenible'.

    Args:
        df (pd.DataFrame): DataFrame con columnas de imputación anual y 'Sostenible'.

    Returns:
        pd.DataFrame: DataFrame con los valores de imputación sostenible actualizados.
    """

    columnas_base = ['Referencia Interna', 'Título', FECHA_INICIO_COL, FECHA_FIN_COL, 'CUANTÍA TOTAL', DURATION_COL, 'Daily Imputation']

    # Años: solo números
    columnas_anhos = sorted([col for col in df.columns if col.isdigit()])

    # Años sostenibles: tipo "2024_sostenible"
    columnas_sostenibles = sorted([col for col in df.columns if col.endswith('_sostenible') and col.split('_')[0].isdigit()])

    # Construimos columnas finales sin duplicar nada
    columnas_finales = []

    # Añadimos columnas base
    columnas_finales += [col for col in columnas_base if col in df.columns]

    if 'Sostenible' in df.columns:
        columnas_finales.append('Sostenible')

   
    # Añadimos los años normales
    columnas_finales += columnas_anhos

    # Añadimos los años sostenibles
    columnas_finales += columnas_sostenibles

    df[columnas_finales].to_excel(path_salida_temp, index=False)
    
    return path_salida_temp

def obtener_tipo_cambio(anio):
    """
    Obtiene el tipo de cambio de EUR a USD para el 31 de diciembre de un año dado.

    Utiliza la API de Fixer.io para obtener el tipo de cambio histórico.

    Args:
        anio (str | int): El año para el cual se desea obtener el tipo de cambio.

    Returns:
        float | None: El tipo de cambio EUR a USD para la fecha especificada,
                      o None si ocurre un error en la solicitud o la respuesta.
    """
    
    url = f"https://data.fixer.io/api/{anio}-12-31"
    params = {        "access_key": "d1e71d969c6f68cbf05213033123001b","base": "EUR", "symbols": "USD"}
    try:
            res = requests.get(url, params=params)
            res.raise_for_status()
            data = res.json()

            # Depuración: Mostrar toda la respuesta
            print(f"Respuesta de la API para {anio}: {data}")

            # Verificar si la respuesta tiene la clave 'rates'
            if "rates" not in data:
                print(f" No se encontró la clave 'rates' en la respuesta para {anio}. Respuesta completa: {data}")
                return None

            # Asegurarse de que la clave 'USD' está en 'rates'
            if "USD" not in data["rates"]:
                print(f" No se encontró 'USD' en 'rates' para {anio}. Respuesta: {data['rates']}")
                return None

            return data["rates"]["USD"]

    except requests.exceptions.RequestException as e:
        print(f"Error en la solicitud HTTP: {e}")
        return None

def generar_tabla_resumen(path_salida_temp, path_salida_final):
    """
    Genera el archivo Excel final con la tabla de resumen financiero.

    Lee el archivo temporal, procesa los datos para calcular totales en USD
    y ratios, y exporta los datos originales, el resumen y los totales
    a diferentes hojas en el archivo final, aplicando estilos.

    Args:
        path_salida_temp (str): Ruta al archivo Excel temporal con los datos procesados.
        path_salida_final (str): Ruta donde se guardará el archivo Excel final.
    """
    df = pd.read_excel(path_salida_temp)
    df_resumen, df_totales, tipos_cambio = procesar_datos(df)
    exportar_a_excel(df, df_resumen, df_totales, path_salida_temp, path_salida_final)


def procesar_datos(df):
    """
    Procesa el DataFrame para calcular los totales de investigación en EUR y USD,
    así como el ratio de investigación sostenible.

    Args:
        df (pd.DataFrame): DataFrame con los datos de proyectos, incluyendo imputación anual
                           y columnas sostenibles.

    Returns:
        tuple[pd.DataFrame, pd.DataFrame, dict]:
            - df_resumen (pd.DataFrame): DataFrame con el resumen anual por moneda.
            - df_totales (pd.DataFrame): DataFrame con los totales generales y el ratio.
            - tipos_cambio (dict): Diccionario con los tipos de cambio utilizados por año.
    """
    idx_sostenible = df.columns.get_loc('Sostenible')
    columnas_despues = df.columns[idx_sostenible + 1:]

    anios_normales = [col for col in columnas_despues if not col.endswith('_sostenible')]
    anios_sostenibles = [col for col in columnas_despues if col.endswith('_sostenible')]

    anios_normales = anios_normales[1:-1] if len(anios_normales) > 2 else anios_normales
    anios_sostenibles = anios_sostenibles[1:-1] if len(anios_sostenibles) > 2 else anios_sostenibles

    fila_total = df.iloc[-1]
    total_research_euros = [float(fila_total[a]) for a in anios_normales]
    sustainable_research_euros = [float(fila_total[a]) for a in anios_sostenibles]

    tipos_cambio = obtener_tipos_cambio(anios_normales)

    total_research_usd = [round(e * tipos_cambio[a], 2) for e, a in zip(total_research_euros, anios_normales)]
    sustainable_research_usd = [round(e * tipos_cambio[a], 2) for e, a in zip(sustainable_research_euros, anios_normales)]

    total_usd = sum(total_research_usd)
    sustainable_usd = sum(sustainable_research_usd)
    ratio = round((sustainable_usd / total_usd) * 100, 1) if total_usd and sustainable_usd else 0

    df_resumen = pd.DataFrame({
        (TOTAL_RESEARCH_LABEL, anho): [f"{eur:,.2f} €", tipos_cambio[anho], f"${usd:,.1f}"]
        for anho, eur, usd in zip(anios_normales, total_research_euros, total_research_usd)
    } | {
        (SUSTAIN_RESEARCH_LABEL, anho.split('_')[0]): [f"{eur:,.2f} €", tipos_cambio[anho.split('_')[0]], f"${usd:,.1f}"]
        for anho, eur, usd in zip(anios_sostenibles, sustainable_research_euros, sustainable_research_usd)
    }, index=CURRENCY_INDEX)

    df_totales = pd.DataFrame({
        "": [
            "Total research funds (in US Dollars)",
            "Total research funds dedicated to sustainability research",
            "The ratio of sustainability research funding to total research funding"
        ],
        "Resultados": [
            f"${total_usd:,.1f}",
            f"${sustainable_usd:,.1f}",
            f"{ratio}%",
        ]
    })

    print(total_usd, sustainable_usd)
    return df_resumen, df_totales, tipos_cambio


def obtener_tipos_cambio(anios):
    """
    Obtiene los tipos de cambio para una lista de años.

    Llama a `obtener_tipo_cambio` para cada año y maneja los casos en que
    no se pueda obtener el tipo de cambio.

    Args:
        anios (list[str]): Lista de años (como strings) para los que se necesitan tipos de cambio.

    Returns:
        dict: Diccionario donde las claves son los años y los valores son los tipos de cambio.
              Usa 0 como valor por defecto si no se puede obtener el tipo de cambio para un año.
    """
    tipos = {}
    for a in anios:
        cambio = obtener_tipo_cambio(a)
        if cambio is None:
            cambio = 0
            print(f"Advertencia: No se obtuvo el tipo de cambio para el año {a}, usando valor por defecto {cambio}")
        tipos[a] = cambio
    return tipos


def exportar_a_excel(df, df_resumen, df_totales, path_temp, path_final):
    """
    Exporta los DataFrames a un archivo Excel con múltiples hojas y aplica estilos.

    Guarda los datos originales, el resumen financiero y los totales en hojas separadas
    y aplica formato condicional y estilos a la hoja de resumen.

    Args:
        df (pd.DataFrame): DataFrame con los datos originales procesados.
        df_resumen (pd.DataFrame): DataFrame con el resumen financiero anual.
        df_totales (pd.DataFrame): DataFrame con los totales generales y el ratio.
        path_temp (str): Ruta al archivo Excel temporal.
        path_final (str): Ruta donde se guardará el archivo Excel final.
    """
    
    with pd.ExcelWriter(path_temp, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name="Datos Originales", index=False)
        df_resumen.to_excel(writer, sheet_name=RESUMEN_FINANCIERO, startrow=0)
        df_totales.to_excel(writer, sheet_name=RESUMEN_FINANCIERO, startrow=len(df_resumen.index) + 4, index=False)

    wb = load_workbook(path_temp)
    ws = wb[RESUMEN_FINANCIERO]
    aplicar_estilos(ws, len(df_resumen.index))
    wb.save(path_final)
    print(f"Tabla de resumen exportada con éxito en el archivo: {path_final}")


def aplicar_estilos(ws, resumen_rows):
    """
    Aplica estilos de formato (colores, fuentes, alineación) a la hoja de resumen financiero.

    Args:
        ws: Objeto Worksheet de openpyxl.
        resumen_rows (int): Número de filas en la tabla de resumen anual (para determinar
                            el inicio de la tabla de totales).
    """
    max_col = ws.max_column
    current_fill = None

    for col in range(2, max_col + 1):
        col_letter = get_column_letter(col)
        cat_cell = ws[f"{col_letter}1"]
        year_cell = ws[f"{col_letter}2"]

        if cat_cell.value:
            if TOTAL_RESEARCH_LABEL in str(cat_cell.value):
                current_fill = PatternFill(start_color=COLOR_TOTAL, end_color=COLOR_TOTAL, fill_type="solid")
                data_fill = PatternFill(start_color=COLOR_TOTAL_DATA, end_color=COLOR_TOTAL_DATA, fill_type="solid")
            elif SUSTAIN_RESEARCH_LABEL in str(cat_cell.value):
                current_fill = PatternFill(start_color=COLOR_SUSTAIN, end_color=COLOR_SUSTAIN, fill_type="solid")
                data_fill = PatternFill(start_color=COLOR_SUSTAIN_DATA, end_color=COLOR_SUSTAIN_DATA, fill_type="solid")
            cat_cell.fill = current_fill

        if current_fill:
            year_cell.fill = current_fill
            year_cell.font = FONT_PINK
            year_cell.alignment = ALIGN_CENTER
            for row in range(4, 7):
                cell = ws[f"{col_letter}{row}"]
                cell.fill = data_fill
                cell.alignment = ALIGN_CENTER

    aplicar_estilos_totales(ws, resumen_rows + 5)


def aplicar_estilos_totales(ws, start_row):
    """
    Aplica estilos de formato a la tabla de totales generales en la hoja de resumen financiero.

    Esta función itera sobre las filas que contienen los totales generales (total USD,
    sostenible USD, y ratio) y aplica colores de fondo, fuente negrita y alineación
    central a las celdas de resultados, basándose en el tipo de total.

    Args:
        ws: Objeto Worksheet de openpyxl. La hoja de cálculo a la que se aplicarán los estilos.
        start_row (int): La fila de inicio de la tabla de totales generales dentro de la hoja.
                         Esta fila es donde comienza la etiqueta "Total research funds (in US Dollars)".
    """
    for row in range(start_row, start_row + 3):
        label_cell = ws[f"A{row}"]
        result_cell = ws[f"B{row}"]
        if label_cell.value:
            if "Total research funds (in US Dollars)" in str(label_cell.value):
                result_cell.fill = PatternFill(start_color=COLOR_TOTAL, end_color=COLOR_TOTAL, fill_type="solid")
            elif "Total research funds dedicated to sustainability research" in str(label_cell.value):
                result_cell.fill = PatternFill(start_color=COLOR_SUSTAIN, end_color=COLOR_SUSTAIN, fill_type="solid")
            result_cell.font = FONT_BOLD
            result_cell.alignment = ALIGN_CENTER
 

def eliminar_tablas_vacias(doc):
    """
    Elimina las tablas vacías (sin contenido) del documento Word.

    Itera sobre las tablas en el documento en orden inverso para evitar problemas
    al eliminar elementos mientras se itera. Para cada tabla, verifica si está vacía
    usando `es_tabla_vacia` y la elimina si es necesario.

    Args:
        doc (Document): Objeto Document de python-docx.
    """
    for table in doc.tables:
        if es_tabla_vacia(table):
            eliminar_tabla(table)


def es_tabla_vacia(table):
    """
    Devuelve True si la tabla no contiene texto en ninguna de sus celdas.

    Recorre cada fila de la tabla y utiliza `fila_contiene_texto` para verificar
    si alguna fila tiene contenido. Si encuentra una fila con texto, la tabla
    no está vacía.

    Args:
        table: Objeto Table de python-docx.

    Returns:
        bool: True si la tabla está vacía, False en caso contrario.
    """
    for row in table.rows:
        if fila_contiene_texto(row):
            return False
    return True


def fila_contiene_texto(row):
    """
    Devuelve True si alguna celda de una fila contiene texto.

    Itera sobre las celdas de la fila y verifica si el texto de la celda,
    después de eliminar espacios iniciales/finales, no está vacío.

    Args:
        row: Objeto Row de python-docx.

    Returns:
        bool: True si la fila contiene texto, False en caso contrario.
    """
    return any(cell.text.strip() for cell in row.cells)


def eliminar_tabla(table):
    """
    Elimina una tabla del documento.

    Accede al elemento XML de la tabla y lo elimina de su elemento padre
    en la estructura del documento.

    Args:
        table: Objeto Table de python-docx a eliminar.
    """
    tbl_element = table._element
    tbl_element.getparent().remove(tbl_element)



def fill_description(doc, ruta_archivo):
    """
    Llena la sección de descripción del documento Word.

    Busca el párrafo que contiene "Description:", extrae el rango de años
    del nombre del archivo de entrada, construye una cadena de texto descriptiva
    y reemplaza el contenido del párrafo encontrado con esta nueva descripción.

    Args:
        doc (Document): Objeto Document de python-docx.
        ruta_archivo (str): Ruta al archivo Excel de entrada, utilizada para
                            extraer el rango de años.
    """
    anio_rango = re.search(r'(\d{4})-(\d{4})', ruta_archivo)

    if anio_rango:
        # Extraemos los años de inicio y fin del rango
        anio_inicio = int(anio_rango.group(1))
        anio_fin = int(anio_rango.group(2))
        
        # Generamos una lista con los años intermedios
        anios_intermedios = list(range(anio_inicio + 1, anio_fin))  # excluyendo los límites
        
        print("Años intermedios:", anios_intermedios)
        # Construcción de la descripción
        description_text = (
            f"The complete list of research projects (years {anio_rango.group(0)}), their amount, their start and end date and their distribution in annual amounts is provided as evidence. "
            f"It also includes whether the project is sustainable research or not. "
            f"For the calculation, the distribution has been made based on the days included in each year and the amount has been determined for the three-year evaluation period ({', '.join(map(str, anios_intermedios))})."
        )

    else:
        print("No se encontró un rango de años.")
        
    
    # Reemplazar el texto en el documento
    for para in doc.paragraphs:
        if "Description:" in para.text:
            para.text = f"Description:\n\n{description_text}"
            break


def set_cell_background(cell, color_hex):
    """
    Aplica color de fondo hexadecimal a una celda de tabla en un documento Word.

    Modifica directamente el elemento XML de la celda para añadir el sombreado
    con el color especificado.

    Args:
        cell: Objeto Cell de python-docx.
        color_hex (str): Código hexadecimal del color (ej. "RRGGBB").
    """
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcpr.append(shd)

   

def insertar_tablas_en_documento(doc, df_resumen, df_datos):
    """
    Inserta las tablas de resumen financiero y datos originales en el documento Word.

    Busca un párrafo marcador ('Additional'), elimina las tablas vacías existentes,
    y luego inserta las tablas de resumen y datos antes del párrafo marcador,
    aplicando estilos básicos y colores a la tabla de resumen.

    Args:
        doc (Document): Objeto Document de python-docx.
        df_resumen (pd.DataFrame): DataFrame con la tabla de resumen financiero.
        df_datos (pd.DataFrame): DataFrame con los datos originales procesados.
    """
    if isinstance(df_resumen.columns, pd.MultiIndex):
        df_resumen.columns = [
            '' if 'Unnamed' in str(col) else ' '.join([str(c) for c in col if 'Unnamed' not in str(c)])
            for col in df_resumen.columns
        ]

    eliminar_tablas_vacias(doc)

    colores_encabezado = {
        "total research": "93c5fd",
        "sustainability": "86efac",
        "ratio": "fde68a",
    }

    colores_filas = {
        "ratio of sustainability": "fde68a",
        "total research funds (in us dollars)": "dbeafe",
        "total research funds dedicated to sustainability research": "dcfce7",
    }

    def agregar_tabla(paragraph, dataframe, pintar=True):
        """Función auxiliar para añadir una tabla después de un párrafo."""

        new_paragraph = paragraph.insert_paragraph_before()
        tabla = doc.add_table(rows=1, cols=len(dataframe.columns))
        tabla.style = 'Table Grid'

        hdr_cells = tabla.rows[0].cells
        columnas_verdes = set()
        columnas_azules = set()

        for j, col in enumerate(dataframe.columns):
            col_str = str(col).replace("\n", " ").strip()
            hdr_cells[j].text = ""
            run = hdr_cells[j].paragraphs[0].add_run(col_str)
            run.bold = True
            if not pintar:
                run.font.size = Pt(9)

            col_str_clean = col_str.lower()
            if "sustainability" in col_str_clean:
                set_cell_background(hdr_cells[j], colores_encabezado["sustainability"])
                columnas_verdes.add(j)
            elif "total research" in col_str_clean:
                set_cell_background(hdr_cells[j], colores_encabezado["total research"])
                columnas_azules.add(j)
            elif pintar and re.search(r'20\d{2}', col_str_clean):
                if "sostenible" not in col_str_clean:
                    set_cell_background(hdr_cells[j], colores_encabezado["total research"])
                    columnas_azules.add(j)
                else:
                    set_cell_background(hdr_cells[j], colores_encabezado["sustainability"])
                    columnas_verdes.add(j)

        for _, row in dataframe.iterrows():
            row_cells = tabla.add_row().cells
            fila_no_vacia = False
            for j, value in enumerate(row):
                val = "" if pd.isna(value) else str(value)
                row_cells[j].text = val
                if val:
                    fila_no_vacia = True
                if pintar:
                    if j in columnas_azules:
                        set_cell_background(row_cells[j], "dbeafe")
                    elif j in columnas_verdes:
                        set_cell_background(row_cells[j], "dcfce7")

            if pintar and fila_no_vacia:
                texto = ' '.join(str(row.iloc[0]).split()).lower()
                for clave, color in colores_filas.items():
                    if clave in texto:
                        for cell in row_cells:
                            if cell.text.strip():
                                set_cell_background(cell, color)
                        break

        new_paragraph._element.addnext(tabla._element)

    for paragraph in doc.paragraphs:
        if 'Additional' in paragraph.text:
            agregar_tabla(paragraph, df_resumen, pintar=True)
            agregar_tabla(paragraph, df_datos, pintar=False)
            break


        
def generar_informe(excel_path, excel):
    """
    Genera el informe Word y PDF a partir de un archivo Excel procesado.

    Carga una plantilla de documento Word, reemplaza el encabezado específico
    del reporte 6.4, llena la sección de descripción, inserta las tablas
    de resumen y datos originales, guarda el documento modificado como .docx,
    y luego intenta convertirlo a .pdf.

    Args:
        excel_path (str): Ruta al archivo Excel final que contiene las hojas
                          "Resumen Financiero" y "Datos Originales".
        excel (str): Ruta al archivo Excel de entrada original (usado para
                     extraer el rango de años para la descripción).
    """
    base_dir = os.path.dirname(__file__)
    plantilla_path = os.path.join(base_dir, 'informe_general.docx')

    if not os.path.exists(plantilla_path) or not os.path.exists(excel_path):
        print(" Plantilla o Excel no encontrado")
        return

    output_filename = "University_Country_6_4_Total Research Funds Dedicated to Sustainability Research (in US Dollars)"
    docx_output = os.path.join(base_dir, f"{output_filename}.docx")
    pdf_output = os.path.join(base_dir, f"{output_filename}.pdf")

    doc = Document(plantilla_path)

    # Reemplazo del texto
    for p in doc.paragraphs:
        if "[6] Education and Research (ED)" in p.text:
            p.text = "[6] Education and Research (ED)\n\n[6.4] Total Research Funds Dedicated to Sustainability Research (in US Dollars)\n"

    fill_description(doc, excel)
    # Leer datos
    df_resumen = pd.read_excel(excel_path, sheet_name='Resumen Financiero', header=[0, 1])
    df_datos = pd.read_excel(excel_path, sheet_name='Datos Originales')

    # Insertar tablas
    insertar_tablas_en_documento(doc, df_resumen, df_datos)

    # Guardar y convertir
    doc.save(docx_output)
    print(f"Informe Word guardado: {docx_output}")

    try:
        convert(docx_output, pdf_output)
        print(f"Informe PDF guardado: {pdf_output}")
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")



def generar(excel):
    """
    Función principal para generar el informe 6.4.

    Carga los datos del archivo Excel, los procesa (preparación de fechas,
    cálculo de imputación diaria y anual, marcado de sostenibilidad),
    añade la fila de totales, exporta los resultados a un archivo temporal,
    genera la tabla de resumen financiero y finalmente crea el informe
    en formato Word y PDF.

    Args:
        excel (str): Ruta al archivo Excel de entrada con los datos de proyectos.
    """
    columnas_esperadas = [FECHA_INICIO_COL, FECHA_FIN_COL, 'CUANTÍA TOTAL']
    df = cargar_excel_con_header_dinamico(excel, columnas_esperadas)
    df = preparar_fechas(df)
    df = calcular_imputacion_diaria(df)
    df = imputar_por_anho(df)
    df = marcar_sostenibles_api(df)
    df = duplicar_valores_sostenibles(df)
    df = anhadir_fila_total(df)
    path_salida_temp = exportar_resultado(df, "resultados_finales.xlsx")
    path_salida_final = "archivo_final.xlsx"
    generar_tabla_resumen(path_salida_temp, path_salida_final)
    os.remove(path_salida_temp)
    generar_informe(path_salida_final, excel)
    

    




    

