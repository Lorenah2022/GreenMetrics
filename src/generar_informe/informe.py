import pandas as pd
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os
import sys

def generar_informe(anho, informe):
    # Rutas de los archivos
    base_dir = os.path.dirname(__file__)  # Directorio base donde se encuentra el script
    template_path =  os.path.join(base_dir, 'modelo_base.docx')
    
    # Datos de la Universidad de Burgos
    university_name = "University of Burgos"
    country_name = "Spain"
    web_address = "https://www.ubu.es/"

    if informe == "6_1":
        excel_data_path = combinar_excels()  # Solo combinar si informe == "6_1"
        output_filename = f"{university_name}_{country_name}_6_1_Number_of_courses_or_modules_related_to_environment_and_sustainability_offered"
        headers_custom = ["Course Title", "Degree", "Degree link", "Notes", "Basic Link"]
    else:
        excel_data_path = os.path.join(base_dir, 'data.xlsx')
        output_filename = "output_filled"
        headers_custom = None
        
    output_docx_path =  os.path.join(base_dir, f"{output_filename}.docx")
    output_pdf_path =  os.path.join(base_dir, f"{output_filename}.pdf")
    excel_path = os.path.join('generar_informe', 'courses.xlsx')
    logo_path = os.path.join(base_dir, '../static/images/logo-UBU.jpg') 

   # Leer datos del Excel
    def extract_data_from_excel(excel_path):
        # Leer los datos del archivo Excel
        df = pd.read_excel(excel_path)
        # Solo aplicar el filtro si el informe es "6_1"
        if informe == "6_1":
            # Filtrar las filas que tengan competencias
            df = df[df['Competences'].notna()]
        year = anho
        num_courses = df.shape[0]
        headers = df.columns.tolist() if headers_custom is None else headers_custom
        data = df.values.tolist()
        
        return headers, data, year, num_courses
    
    
    # Es necesaria la conversión de centimetros a pulgadas ya que la biblioteca "pyhton-docx" usa pulgadas.
    def cm_to_inches(cm):
        return cm / 2.54


    # Función para insertar el logo en lugar del texto "Put your university logo here"
    def replace_logo_placeholder(doc, logo_path):
        header = doc.sections[0].header  # Acceder a la cabecera de la primera sección

        for para in header.paragraphs:
            if "Put your university logo here" in para.text:
                inline_shapes = para._element.xpath(".//w:t[text()='Put your university logo here']")
                if inline_shapes:
                    inline_shapes[0].getparent().remove(inline_shapes[0])  # Elimina solo el texto del marcador
                    run = para.add_run()
                    run.add_picture(logo_path, width=Inches(cm_to_inches(3.06)), height=Inches(cm_to_inches(2.25)))  # Ajustado a centímetros
                break
            
    # Función para rellenar la información de la universidad
    def fill_university_info(doc, university_name, country_name, web_address):
        for para in doc.paragraphs:
            if "University" in para.text:
                para.text = f"University : {university_name}"
            elif "Country" in para.text:
                para.text = f"Country : {country_name}"
            elif "Web Address" in para.text:
                para.text = f"Web Address : {web_address}"
    # Función para rellenar la tabla
    def fill_table(doc, headers, data):
        for table in doc.tables:
            # Limpiar la tabla existente
            for row in table.rows:
                for cell in row.cells:
                    cell.text = ""
            
            # Ajustar el número de columnas
            num_columns = len(headers)
            while len(table.columns) < num_columns:
                table.add_column(width=Inches(1.5))
            
            # Insertar encabezados
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            # Insertar datos
            for row_idx, row in enumerate(data):
                if row_idx + 1 >= len(table.rows):
                    table.add_row()
                for col_idx, value in enumerate(row):
                    table.cell(row_idx + 1, col_idx).text = str(value)
            break

    # Función para agregar la descripción
    def fill_description(doc, year, num_courses):
        description_text = (f"The list above includes courses directly related to sustainability offered by the University in the {year} academic year.\n\n"
                            f"Total number of courses with sustainability embedded for courses running in {year}: {num_courses}")
        for para in doc.paragraphs:
            if "Description:" in para.text:
                para.text = f"Description:\n\n{description_text}"
                break

    # Leer documento Word
    doc = Document(template_path)
    # Llamar a la función para reemplazar el marcador del logo
    replace_logo_placeholder(doc, logo_path)

    # Llamar a la función para rellenar la información de la universidad
    fill_university_info(doc, university_name, country_name, web_address)
    headers, data_from_excel, year, num_courses = extract_data_from_excel(excel_data_path)
    fill_table(doc, headers, data_from_excel)
    fill_description(doc, year, num_courses)
    
    # Guardar el documento modificado
    doc.save(output_docx_path)

    # Convertir el documento a PDF
    convert(output_docx_path, output_pdf_path)

    print(f"Documento PDF generado en: {output_pdf_path}")

def combinar_excels():
    base_dir = os.path.dirname(__file__)  
    # Rutas de los archivos
    excel_courses = os.path.join(base_dir, '../sostenibilidad/data/datos_asignaturas_masteres.xlsx')
    excel_guias = os.path.join(base_dir, '../sostenibilidad/data/resultados_guias.xlsx')
    
    # Leer los archivos Excel
    df_courses = pd.read_excel(excel_courses)  # Contiene 'basic_link' y códigos de asignatura
    df_guias = pd.read_excel(excel_guias)  # Contiene los cursos, donde añadiremos la columna 'Link'
    
    # Mostrar las primeras filas para depuración
    print("Primeras filas de df_courses (datos_asignaturas_masteres):")
    print(df_courses.head())
    print("\nPrimeras filas de df_guias (resultados_guias):")
    print(df_guias.head())
    
    # Asegurarse de que las columnas de código de asignatura coinciden
    # Mapear la columna 'Code' de df_guias con la columna 'codigo_asignatura' de df_courses para obtener 'basic_link'
    df_guias['Link'] = df_guias['Code'].map(df_courses.set_index('codigo_asignatura')['basic_link'])
    
    # Verifica que la columna Link se haya añadido correctamente
    print("\nPrimeras filas después de añadir la columna Link:")
    print(df_guias.head())
    
    # Guardar el resultado combinado en un nuevo archivo, asegurando que no se pierdan otras columnas
    combined_path = os.path.join(base_dir, 'resultados_guias_con_link.xlsx')
    df_guias.to_excel(combined_path, index=False)
    
    return combined_path

if __name__ == "__main__":
    # Obtener el año académico desde los argumentos de la línea de comandos
    if len(sys.argv) != 3:
        print("Uso correcto: python script.py <anho> <informe>")
        sys.exit(1)
        
    anho = sys.argv[1]
    informe = sys.argv[2]
    generar_informe(anho, informe)