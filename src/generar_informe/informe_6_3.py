import sqlite3
import os
import subprocess
import sys
from docx import Document
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.table import Table
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table


SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# Agregar `src` al sys.path 
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)

def verificar_tipos_programa(anio, db_path):
    """
    Verifica si existen datos para los tipos de programa 'grado' y 'master' en un año dado.
    Además, cuenta cuántos programas son sostenibles ('sí' en la columna 'sostenible').

    Parámetros:
    - anio (int o str): El año que se desea verificar.
    - db_path (str): Ruta a la base de datos SQLite.

    Retorna:
    - dict: Información de existencia, cantidad total y cantidad sostenible para cada tipo.
            Ejemplo: {
                'grado': {'existe': True, 'cantidad_total': 100, 'cantidad_sostenible': 50},
                'master': {'existe': False, 'cantidad_total': 0, 'cantidad_sostenible': 0}
            }
    """
    if not os.path.exists(db_path):
        raise FileNotFoundError(f"No se encontró la base de datos en {db_path}")

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    tipos = ['grado', 'master']
    resultado = {}

    try:
        for tipo in tipos:
            # Contar total
            cursor.execute("""
                SELECT COUNT(*)
                FROM busqueda
                WHERE anho = ? AND tipo_programa = ?
            """, (anio, tipo))
            cantidad_total = cursor.fetchone()[0]

            # Contar sostenibles
            cursor.execute("""
                SELECT COUNT(*)
                FROM busqueda
                WHERE anho = ? AND tipo_programa = ? AND TRIM(sostenibilidad) = 'Sí'
            """, (anio, tipo))
            cantidad_sostenible = cursor.fetchone()[0]

            resultado[tipo] = {
                'existe': cantidad_total > 0,
                'cantidad_total': cantidad_total,
                'cantidad_sostenible': cantidad_sostenible
            }

        return resultado


    finally:
        conn.close()



def add_hyperlink(paragraph, url, text):
    """
    Agrega un hipervínculo a un párrafo en un documento Word.

    Parámetros:
    - paragraph (Paragraph): El objeto Paragraph al que se añadirá el hipervínculo.
    - url (str): La URL del hipervínculo.
    - text (str): El texto visible del hipervínculo.
    """
    # Limpiar espacios no rompibles y espacios finales/iniciales
    url = url.replace("\u00A0", "").strip()

    # Crear la relación del hipervínculo en el documento
    part = paragraph._element
    rid = paragraph._parent.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)

    # Crear el run del enlace
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # Color azul
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # Azul
    rpr.append(color)

    # Subrayado manual
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")  # Subrayado
    rpr.append(u)

    new_run.append(rpr)

    # Agregar el texto visible
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    part.append(hyperlink)


def insert_table_after_paragraph(paragraph: Paragraph, doc: Document, rows: int, cols: int) -> Table:
    """
    Inserta una tabla justo después de un párrafo dado y con estilo visual.

    Parámetros:
    - paragraph (Paragraph): El párrafo después del cual se insertará la tabla.
    - doc (Document): El objeto Document de python-docx.
    - rows (int): Número de filas para la nueva tabla.
    - cols (int): Número de columnas para la nueva tabla.

    Retorna:
    - Table: El objeto Table recién creado.
    """
    # Insertar un párrafo nuevo justo después del original
    new_para_element = OxmlElement("w:p")
    paragraph._element.addnext(new_para_element)
    new_paragraph = Paragraph(new_para_element, paragraph._parent)

    # Insertar la tabla en ese nuevo párrafo
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'  # Asegura líneas visibles

    # Mover la tabla al lugar correcto
    tbl_element = table._element
    new_paragraph._element.addnext(tbl_element)

    return table

def fill_description(doc, anho, total, total_sustainable):
    """
    Llena la sección de descripción del documento con texto, una tabla y texto adicional.

    Busca el párrafo que contiene "Description:" y lo reemplaza, insertando
    una tabla con el año, total de cursos sostenibles y total de cursos,
    seguida de un párrafo con el ratio calculado.

    Parámetros:
    - doc (Document): El objeto Document de python-docx.
    - anho (str): El año a incluir en la descripción y tabla.
    - total (int): El número total de cursos a incluir en la descripción y tabla.
    - total_sustainable (int): El número total de cursos sostenibles a incluir.
    """
    description_text = (
        "Ratio of sustainability courses to total courses/subjects"
    )

    for para in doc.paragraphs:
        if "Description:" in para.text:
            # Reemplazar el texto con el texto de introducción
            para.clear()

            # Añadimos el texto "Description:" en negrita
            run_desc = para.add_run("Description:\n\n")
            run_desc.bold = True

            # Añadimos el texto descriptivo en cursiva
            run_italic = para.add_run(description_text)
            run_italic.italic = True

            # Insertar tabla justo después de este párrafo
            table = insert_table_after_paragraph(para, doc, 2, 3)

            # Llenar tabla
            table.cell(0, 0).text = "Year"
            table.cell(0, 1).text = "Total Sustainable Courses"
            table.cell(0, 2).text = "Total Courses"

            
            table.cell(1, 0).text = anho
            table.cell(1, 1).text = str(total_sustainable)
            table.cell(1, 2).text = str(total)

            # Calcular el ratio de sostenibilidad
            if total != 0:  # Para evitar la división por cero
                ratio = total_sustainable / total
            else:
                ratio = 0  
                

            tbl_element = table._element
            new_para_element = OxmlElement("w:p")
            tbl_element.addnext(new_para_element)
            new_paragraph = Paragraph(new_para_element, para._parent)
            new_paragraph.add_run(f"Ratio: {ratio * 100:.2f}%")


            break

def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str = None) -> Paragraph:
    """
    Inserta un nuevo párrafo justo después del párrafo dado.

    Parámetros:
    - paragraph (Paragraph): El párrafo después del cual se insertará el nuevo párrafo.
    - text (str, optional): El texto inicial para el nuevo párrafo. Por defecto es "".
    - style (str, optional): El estilo para el nuevo párrafo. Por defecto es None.

    Retorna:
    - Paragraph: El objeto Paragraph recién creado.
    """
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style
    return new_paragraph

   

def generar_informe(total_general, anio, sostenible_total):
    """
    Genera el informe Word y PDF para el reporte 6.3.

    Utiliza una plantilla .docx, reemplaza el encabezado, añade hipervínculos,
    llena la sección de descripción con una tabla y el ratio calculado,
    y guarda el resultado en formato .docx y .pdf.

    Parámetros:
    - total_general (int): El número total de cursos/materias a incluir en el informe.
    - anio (str): El año para el cual se genera el informe.
    - sostenible_total (int): El número total de cursos sostenibles a incluir.
    """
    
    base_dir = os.path.dirname(__file__)  # Directorio base donde se encuentra el script
    template_path = os.path.join(base_dir, 'informe_general.docx')
    
    if not os.path.exists(template_path):
        print(f"❌ Error: No se encontró la plantilla en {template_path}")
        sys.exit(1)

    output_filename = "University_Country_6_3_Ratio_of_sustainability_courses_to_total_courses"
    
    doc = Document(template_path)

    # 1. Eliminar cualquier tabla vacía en el documento
    eliminar_tablas_vacias(doc)

    # 2. Reemplazar encabezado
    for paragraph in doc.paragraphs:
        if "[6] Education and Research (ED)" in paragraph.text:
            paragraph.text = "[6] Education and Research (ED)\n\n[6.3]  Ratio of sustainability courses to total courses/subjects\n\n"          
            # Agregar el primer enlace debajo de la sección "[6.2]"
            p1 =insert_paragraph_after(paragraph)

            enlace1 = "https://www.ubu.es/grados-ordenados-por-ramas-de-conocimiento"
            add_hyperlink(p1, enlace1, enlace1)

            # Agregar el segundo enlace debajo del primero
            p2 = insert_paragraph_after(paragraph)

            enlace2 = "https://www.ubu.es/estudios/oferta-de-estudios/masteres-universitarios-oficiales"
            add_hyperlink(p2, enlace2, enlace2)
            
            break
        

    # 4. Añadir la descripción con tabla y datos
    fill_description(doc, anio, total_general, sostenible_total)

    report_dir = os.path.join(SRC_DIR, "generated_reports", "report_6_3", anio)
    os.makedirs(report_dir, exist_ok=True)  # Crea la carpeta si no existe

    output_docx_path = os.path.join(report_dir, f"{output_filename}.docx")
    output_pdf_path = os.path.join(report_dir, f"{output_filename}.pdf")

    # 5. Guardar en Word
    doc.save(output_docx_path)

    # 6. Intentar conversión a PDF
    try:
        convert(output_docx_path, output_pdf_path)
        print(f"\nDocumento PDF generado en: {output_pdf_path}")
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

 
 

def eliminar_tablas_vacias(doc):
    """
    Elimina las tablas vacías (sin contenido) del documento Word.

    Parámetros:
    - doc (Document): El objeto Document de python-docx.
    """
    for table in doc.tables:
        if es_tabla_vacia(table):
            eliminar_tabla(table)


def es_tabla_vacia(table):
    """
    Verifica si una tabla no contiene texto en ninguna de sus celdas.

    Parámetros:
    - table (Table): El objeto Table de python-docx.

    Retorna:
    - bool: True si la tabla está vacía, False en caso contrario.
    """
    for row in table.rows:
        if fila_contiene_texto(row):
            return False
    return True


def fila_contiene_texto(row):
    """
    Verifica si alguna celda de una fila contiene texto.

    Parámetros:
    - row: El objeto Row de python-docx.

    Retorna:
    - bool: True si la fila contiene texto, False en caso contrario.
    """
    return any(cell.text.strip() for cell in row.cells)


def eliminar_tabla(table):
    """
    Elimina una tabla del documento.

    Parámetros:
    - table (Table): El objeto Table de python-docx a eliminar.
    """
    tbl_element = table._element
    tbl_element.getparent().remove(tbl_element)



def generar(anio):
    """
    Función principal para iniciar la generación del informe 6.3.

    Verifica la existencia de datos de 'grado' y 'master' para el año especificado
    en la base de datos, incluyendo el conteo de cursos sostenibles.
    Si los datos están completos, calcula los totales y genera el informe Word y PDF.
    Si faltan datos, informa al usuario.

    Parámetros:
    - anio (str): El año para el cual se desea generar el informe.
    """
    
    base_dir = os.path.dirname(__file__)  # carpeta actual (\src)
    db_path = os.path.join(base_dir, '..', 'instance', 'busqueda.db')
    db_path = os.path.abspath(db_path)    
  
    resultados = verificar_tipos_programa(anio, db_path)

    if all(info['existe'] for info in resultados.values()):
        print("\nDatos completos. Generando informe...")

        total_grado = resultados['grado']['cantidad_total']
        total_master = resultados['master']['cantidad_total']
        total_general = total_grado + total_master

        sostenible_grado = resultados['grado']['cantidad_sostenible']
        sostenible_master = resultados['master']['cantidad_sostenible']
        sostenible_total = sostenible_grado + sostenible_master

        print(f"Total de cursos: {total_general}")
        print(f"Cursos sostenibles: {sostenible_total}")

        # Aquí podrías pasar también los datos de sostenibles a tu informe si quieres.
        generar_informe(total_general, anio,sostenible_total)

            
    else:
        for tipo, info in resultados.items():
            if not info['existe']:
                print(f"- No hay datos para '{tipo}'. Descargue los datos faltantes.")



