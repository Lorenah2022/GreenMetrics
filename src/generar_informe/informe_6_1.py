import pandas as pd
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys

from helpers import add_hyperlink

SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Agregar `src` al sys.path 
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)

def extract_data_from_excel(path, headers_custom, anho):
    """
        Extrae datos de un archivo Excel, filtrando por la columna 'Competences'.

        Args:
            path (str): La ruta al archivo Excel.
            headers_custom (list, optional): Lista de encabezados personalizados. Si None, se usan los encabezados del archivo.
            anho (str): El año proporcionado a la función generar.

        Returns:
            tuple: Una tupla que contiene:
                   - list: Encabezados de las columnas.
                   - list: Datos extraídos como lista de listas.
                   - str: El año proporcionado a la función generar.
                   - int: El número de filas después del filtrado.
    """
    df = pd.read_excel(path)
    df = df[df['Competences'].notna()]
    headers = df.columns.tolist() if headers_custom is None else headers_custom
    return headers, df.values.tolist(), anho, df.shape[0]


def fill_table(doc, headers, data):
    """
        Llena la primera tabla encontrada en un documento Word con los datos proporcionados.

        Limpia la tabla existente, ajusta el número de columnas, inserta encabezados
        y luego inserta los datos, manejando la creación de hipervínculos para URLs.

        Args:
            doc (docx.document.Document): El objeto documento de python-docx.
            headers (list): Una lista de cadenas para usar como encabezados de la tabla.
            data (list): Una lista de listas, donde cada sublista representa una fila de datos.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = ""
        
        num_columns = len(headers)
        while len(table.columns) < num_columns:
            table.add_column(width=Inches(1.5))
        
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        for row_idx, row in enumerate(data):
            if row_idx + 1 >= len(table.rows):
                table.add_row()
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx + 1, col_idx)
                if col_idx == 2 and isinstance(value, str) and value.startswith("http"):
                    cell.text = ""
                    add_hyperlink(cell.paragraphs[0], value, "Link")
                else:
                    cell.text = str(value)
        break


def fill_description(doc, year, num_courses):
    """
        Rellena el párrafo de descripción en el documento Word con información sobre el año
        y el número total de cursos.

        Busca un párrafo que contenga "Description:" y lo reemplaza con el texto
        de descripción formateado.

        Args:
            doc (docx.document.Document): El objeto documento de python-docx.
            year (str): El año académico.
            num_courses (int): El número total de cursos.
    """
    description_text = (
        f"The list above includes courses directly related to sustainability offered by the University in the {year} academic year.\n\n"
        f"Total number of courses with sustainability embedded for courses running in {year}: {num_courses}"
    )
    for para in doc.paragraphs:
        if "Description:" in para.text:
            para.text = f"Description:\n\n{description_text}"
            break


def generar(anho):
    """
    Genera el informe 6.1 sobre cursos relacionados con medio ambiente y sostenibilidad.

    Combina datos de dos archivos Excel, extrae la información relevante,
    llena una plantilla de documento Word con los datos y una descripción,
    y guarda el resultado como archivo Word y PDF.

    Args:
        anho (str): El año académico para el cual se genera el informe.
    """
    base_dir = os.path.dirname(__file__)
    template_path = os.path.join(base_dir, 'informe_general.docx')

    if not os.path.exists(template_path):
        print(f"Error: No se encontró la plantilla {template_path}")
        sys.exit(1)

    excel_data_path = combinar_excels()
    output_filename = (
        "University_Country_6_1_Number_of_courses_or_modules_related_to_environment_and_sustainability_offered"
    )
    headers_custom = ["Course Title", "Degree", "Degree link", "Notes"]

    if not os.path.exists(excel_data_path):
        print(f"Error: No se encontró el archivo de datos {excel_data_path}")
        sys.exit(1)

    report_dir = os.path.join(SRC_DIR, "generated_reports", "report_6_1", anho)
    os.makedirs(report_dir, exist_ok=True)

    output_docx_path = os.path.join(report_dir, f"{output_filename}.docx")
    output_pdf_path = os.path.join(report_dir, f"{output_filename}.pdf")

    doc = Document(template_path)

    headers, data_from_excel, year, num_courses = extract_data_from_excel(
        excel_data_path, headers_custom, anho
    )
    fill_table(doc, headers, data_from_excel)
    fill_description(doc, year, num_courses)
    doc.save(output_docx_path)

    try:
        convert(output_docx_path, output_pdf_path)
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

    print(f"Documento PDF generado en: {output_pdf_path}")
    
def combinar_excels():
    """
    Combina datos de dos archivos Excel ('datos_asignaturas_masteres.xlsx' y 'resultados_guias.xlsx').

    Realiza un merge basado en códigos de asignatura, renombra una columna,
    filtra filas donde la columna 'Competences' no está vacía, selecciona
    columnas específicas y guarda el resultado en un nuevo archivo Excel
    llamado 'resultados_guias_con_link.xlsx'.

    Returns:
        str: La ruta absoluta al archivo Excel combinado generado.
    """
    base_dir = os.path.dirname(__file__)
    excel_courses = os.path.join(base_dir, '../sostenibilidad/data/datos_asignaturas_masteres.xlsx')
    excel_guias = os.path.join(base_dir, '../sostenibilidad/data/resultados_guias.xlsx')

    if not os.path.exists(excel_courses) or not os.path.exists(excel_guias):
        print("Error: No se encontraron los archivos de datos necesarios.")
        sys.exit(1)

    # Leer los archivos de Excel
    df_courses = pd.read_excel(excel_courses)
    df_guias = pd.read_excel(excel_guias)

    # Asegurar que los valores de "Code" y "codigo_asignatura" sean del mismo tipo
    df_guias["Code"] = df_guias["Code"].astype(str).str.strip()
    df_courses["codigo_asignatura"] = df_courses["codigo_asignatura"].astype(str).str.strip()

    # Realizar el merge basado en la clave correcta
    df_combined = df_guias.merge(df_courses[['codigo_asignatura', 'basic_link']], 
                                 left_on="Code", 
                                 right_on="codigo_asignatura", 
                                 how="right")

    # Renombrar la columna "basic_link" a "Link"
    df_combined.rename(columns={"basic_link": "Link"}, inplace=True)
    
    #Filtrar solo filas donde "Competences" no esté vacío
    df_combined = df_combined[df_combined["Competences"].notna() & (df_combined["Competences"].str.strip() != "")]

    # Seleccionar solo las columnas necesarias
    columnas_deseadas = ["Name", "Degree_Master", "Link","Competences"]
    df_combined = df_combined[columnas_deseadas]

    # Guardar el nuevo Excel con solo las columnas seleccionadas
    combined_path = os.path.join(base_dir, 'resultados_guias_con_link.xlsx')   
    df_combined.to_excel(combined_path, index=False)
    
    return combined_path



