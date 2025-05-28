import csv
import io
import json
import os
import sys
import time
import tkinter as tk
from tkinter import messagebox, simpledialog, ttk

import pandas as pd
import requests
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
from fpdf import FPDF
from PyPDF2 import PdfReader
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException

from deep_translator import GoogleTranslator

# Obtener la ruta absoluta del directorio `src`
SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Agregar `src` al sys.path 
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)

from config import cargar_configuracion


# Variable global
base_dir = os.path.dirname(__file__)
load_dotenv()
# Configuración de la API
config = cargar_configuracion()
base_url = config["base_url"]
api_key = config["api_key"]
myModel = config["model"]

# Función para reemplazar texto en un documento de Word
def replace_text_in_docx(doc, old_text, new_text):
    """
    Reemplaza todas las ocurrencias de un texto antiguo con un texto nuevo en un documento de Word.

    Args:
        doc (docx.document.Document): El objeto documento de python-docx.
        old_text (str): El texto a buscar y reemplazar.
        new_text (str): El texto con el que se reemplazará.
    """
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = new_text

def remove_text_from_docx(doc, text_to_remove):
    """
    Elimina cualquier párrafo que contenga el texto especificado en un documento de Word.

    Args:
        doc (docx.document.Document): El objeto documento de python-docx.
        text_to_remove (str): El texto a buscar en los párrafos para eliminarlos.
    """
    for paragraph in doc.paragraphs:
        if text_to_remove in paragraph.text:
            paragraph.text = ""  # Eliminar el texto



def obtener_html(url):
    """
    Obtiene el contenido HTML de una URL dada.

    Args:
        url (str): La URL de la página web.

    Returns:
        str or None: El contenido HTML de la página si la solicitud fue exitosa (código 200),
                     de lo contrario, None.
    """
    try:
        response = requests.get(url, timeout=10)
        return response.text if response.status_code == 200 else None
    except requests.exceptions.RequestException:
        return None
        
def limpiar_html(html_content):
    """
    Limpia el contenido HTML extrayendo solo el texto y limitándolo a un número de palabras.

    Args:
        html_content (str): El contenido HTML a limpiar.

    Returns:
        str: El texto limpio y truncado.
    """
    soup = BeautifulSoup(html_content, "html.parser")
    texto_limpio = soup.get_text(separator="\n", strip=True)
    return " ".join(texto_limpio.split()[:3500])
    
def extraer_datos_llm(html_texto, url):
    """
    Extrae datos estructurados de texto HTML utilizando un modelo de lenguaje grande (LLM).

    Envía el texto HTML a la API del LLM con instrucciones específicas para extraer
    campos como Edificio, Contrato, Tipo de Mantenimiento, Expediente y Enlace.

    Args:
        html_texto (str): El texto limpio extraído del HTML.
        url (str): La URL de origen del HTML.

    Returns:
        dict or None: Un diccionario con los datos extraídos si la API responde correctamente
                      y los datos tienen el formato esperado, de lo contrario, None.
    """
    if not html_texto:
        return None

    body = {
        "model": myModel,
        "messages": [
            {
                "role": "system",
                "content": (
                    "Extract and return only the following fields:\n"
                    "1. Building: Extracted from Objeto. If the text mentions 'buildings', return 'All university buildings'.\n"
                    "2. Contract: Extracted from the CPV section, including the text after the number [CPV Code]. This will be the detailed description of the contract, e.g., 'Servicios de reparación y mantenimiento de equipos eléctricos de edificios'.\n"
                    "3. Maintenance Type:  Return the information following the format X(Y), where X is the type and Y the number\n"
                        "Base on the type of contract chose to which one it belongs."
                        "1. Preventive Maintenance --- Routine maintenance tasks performed to prevent equipment failures and extend the lifespan of building systems\n"
                        "2. Corrective Maintenance --- Reactive maintenance tasks performed to correct issues as they arise\n"
                        "3. Predictive Maintenance --- Maintenance activities based on the analysis of data and condition monitoring to predict and prevent potential failures \n"
                        "4. Routine Maintenance --- Regular, often daily or weekly, maintenance tasks that ensure the smooth operation and cleanliness of campus buildings\n"
                        "5. Emergency Maintenance --- Urgent maintenance tasks performed in response to unexpected breakdowns or safety hazards that require immediate attention\n"
                        "6. Deferred Maintenance --- Maintenance tasks that are postponed due to budget constraints, resource limitations, or scheduling issues\n"
                        "7. Sustainable Maintenance --- Maintenance activities focused on sustainability and energy efficiency to reduce environmental impact\n"
                        "8. Capital Maintenance --- Large-scale maintenance projects that involve significant investments and are often planned and budgeted for in advance \n"
                        "9. Seasonal Maintenance --- Maintenance tasks specific to certain times of the year to prepare buildings for seasonal changes \n"
                        "10. Compliance Maintenance --- Maintenance activities conducted to ensure compliance with legal, safety, and regulatory standards\n"
                        "11. Custodial Maintenance --- Daily cleaning and janitorial tasks that maintain the cleanliness and hygiene of campus buildings\n"
                        "12. Technical Maintenance --- Specialized maintenance tasks that require technical knowledge and skills\n"
                        "13. Grounds Maintenance --- Maintenance tasks focused on the outdoor areas and landscaping of the campus \n"
                        "14. Building Services Maintenance --- Maintenance of essential building services and utilities \n"

                    "4. File: Extracted from 'Expediente'. If the text appears in the format 'Expediente X UBU/2023/0018 (Company Name)', return only 'UBU/2023/0018' and ignore the company name inside parentheses.\n"
                    "5. Link:"
                    "Extract and return only the following fields in a single line, separated by commas in this format:\n"
                    "Building,Contract,Maintenance Type,File,Link\n"
                    "Do NOT include extra text (like the format), explanations, or headers.\n"
                    "Ensure that you return ONLY these five fields in a single line, without extra text."

                        )
            },
            {
                "role": "user",
                "content": html_texto
            }
        ],
        "temperature": 0.2
    }

    body_json = json.dumps(body)

    response = requests.post(
        f"{base_url}/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        data=body_json
    )

    if response.status_code != 200:
        print(f"Error en la API para {url}: {response.status_code} - {response.text}")
        return None


    #  Obtener respuesta limpia
    message_content = response.json().get('choices', [{}])[0].get('message', {}).get('content', '').strip()

    #  Eliminar `</think>` si existe
    if "</think>" in message_content:
        message_content = message_content.split("</think>")[-1].strip()

    #  Procesar como CSV usando `csv.reader`
    try:
        csv_reader = csv.reader(io.StringIO(message_content))  # Convertir texto a archivo virtual
        contenido_separado = next(csv_reader)  # Obtener la primera (y única) línea

    except Exception as e:
        print(f"Error al leer CSV: {e}\nTexto recibido:\n{message_content}")
        return None

    #  Validar si tiene los 5 elementos requeridos
    if len(contenido_separado) >= 5:
        return {
            "Building": contenido_separado[0].strip(),
            "Contract": contenido_separado[1].strip(),
            "Maintenance Type": contenido_separado[2].strip(),
            "File": contenido_separado[3].strip(),
            "Link": url
        }
    else:
        print(f" Error: La respuesta de la IA no tiene el formato correcto.\nRecibido:\n{message_content}")
        return None
        
def ejecutar_API(enlaces):
    """
    Ejecuta la extracción de datos utilizando la API del LLM para una lista de enlaces.

    Obtiene el HTML de cada enlace, lo limpia, extrae los datos con el LLM,
    filtra los resultados nulos y elimina duplicados basados en el campo "File".

    Args:
        enlaces (list): Una lista de URLs para procesar.

    Returns:
        list: Una lista de diccionarios, donde cada diccionario contiene los datos
              extraídos para un enlace único.
    """
    # Paso 1: Obtener el HTML de los enlaces y extraer los datos
    resultados = []
    for enlace in enlaces:
        html = obtener_html(enlace)
        if html:  # Verificar si el HTML fue obtenido correctamente
            datos = extraer_datos_llm(limpiar_html(html), enlace)
            resultados.append(datos)

    # Paso 2: Filtrar los resultados que no sean None
    resultados_filtrados = []  # Inicializamos una lista vacía para almacenar los resultados filtrados

    for res in resultados:  # Iteramos sobre todos los elementos de la lista "resultados"
        if res:  # Verificamos si el elemento "res" es considerado verdadero (es decir, no es None, vacío, etc.)
            resultados_filtrados.append(res)  # Si es verdadero, lo agregamos a la lista "resultados_filtrados"
    print(resultados)
    # Paso 3: Eliminar duplicados basados en el valor de "File"
    resultados_unicos = {}
    for res in resultados_filtrados:
        if res and "File" in res:
            resultados_unicos[res["File"]] = res

    # Convertir el diccionario de resultados únicos a una lista de valores
    resultados_filtrados_unicos = list(resultados_unicos.values())
    print("Resultados de los enlaces", resultados_filtrados_unicos)
    return list(resultados_filtrados_unicos)

def initialize_table(doc, headers):
    """
    Inicializa la primera tabla encontrada en el documento de Word con los encabezados especificados.

    Asegura que la tabla tenga suficientes columnas para los encabezados y establece el texto
    de la primera fila con los nombres de los encabezados.

    Args:
        doc (docx.document.Document): El objeto documento de python-docx.
        headers (list): Una lista de cadenas para usar como encabezados de la tabla.

    Returns:
        docx.table._Table or None: El objeto tabla inicializado si se encuentra una tabla,
                                   de lo contrario, None.
    """
    for table in doc.tables:
        while len(table.columns) < len(headers):
            table.add_column(width=Inches(1.5))
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        return table  # Retornamos la primera tabla encontrada
    return None

def fill_table(table, headers, data):
    """
    Llena una tabla de Word con datos, incluyendo encabezados y manejo de hipervínculos.

    Elimina las filas existentes (excepto la de encabezados si ya existe) y agrega
    nuevas filas con los datos proporcionados. Si un valor de dato es una URL,
    crea un hipervínculo en la celda.

    Args:
        table (docx.table._Table): El objeto tabla de python-docx a llenar.
        headers (list): Una lista de cadenas que representan los encabezados de la tabla.
        data (list): Una lista de listas o diccionarios, donde cada elemento representa una fila de datos.
    """
    
    # Eliminar todas las filas existentes si la tabla está vacía
    while len(table.rows) > 0:
        table._element.remove(table.rows[0]._element)

    # Crear la fila de encabezados
    header_row = table.add_row()
    for i, header in enumerate(headers):
        header_row.cells[i].text = header

    # Agregar filas con datos
    for row_data in data:
        row = table.add_row()
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            if isinstance(value, str) and (value.startswith("http://") or value.startswith("https://")):
                add_hyperlink(cell.paragraphs[0], value, "Link")
            else:
                cell.text = str(value)

 
# Función que crea un hipervínculo, para el correcto funcionamiento de los enlaces
def add_hyperlink(paragraph, url, text):
    """
    Agrega un hipervínculo a un párrafo en un documento Word.

    Crea la relación del hipervínculo, configura el estilo (color azul y subrayado)
    y añade el texto visible del enlace al párrafo.

    Args:
        paragraph (docx.text.paragraph.Paragraph): El objeto párrafo de python-docx.
        url (str): La URL de destino del hipervínculo.
        text (str): El texto visible del hipervínculo.
    """
    # Crear la relación del hipervínculo en el documento
    part = paragraph._element
    rId = paragraph._parent.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rId)

    # Crear el run del enlace
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Color azul
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # Azul
    rPr.append(color)

    # Subrayado manual
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")  # Subrayado
    rPr.append(u)

    new_run.append(rPr)

    # Agregar el texto visible
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    part.append(hyperlink)
     
# Función que genera el informe
def generar_informe(datos):
    """
    Genera el informe en formato Word y PDF utilizando una plantilla y los datos proporcionados.

    Carga la plantilla 'informe_general.docx', reemplaza texto específico,
    inicializa y llena una tabla con los datos, guarda el documento Word
    y lo convierte a PDF.

    Args:
        datos (list): Una lista de diccionarios, donde cada diccionario contiene
                      los datos para una fila de la tabla del informe.
    """
    base_dir = os.path.dirname(__file__)
    template_path = os.path.join(base_dir, 'informe_general.docx')
    print("Datos recibidos ", datos)
    if not os.path.exists(template_path):
        print(f"Error: No se encontró la plantilla {template_path}")
        return

    output_filename = "University_Country_1_19_Percentage_of_operation_and_maintenance_activities_of_building_in_one_year"
    
    doc = Document(template_path)

    # Reemplazar texto en la plantilla
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace(
            "[6] Education and Research (ED)", 
            "[1] Setting and Infrastructure (SI)\n\n"
            "[1.19] Percentage of operation and maintenance activities of building in one year period\n\n"
            "*Min. at least five maintenance classification for each building"
        )

    # Crear la tabla
    headers = ["Building", "Contract", "Maintenance Type", "File", "Link"]
    table = initialize_table(doc, headers)

    fill_table(table, headers, [list(d.values()) for d in datos])

    report_dir = os.path.join(SRC_DIR, "generated_reports", "report_1_19")
    os.makedirs(report_dir, exist_ok=True)  # Crea la carpeta si no existe

    output_docx_path = os.path.join(report_dir, f"{output_filename}.docx")
    output_pdf_path = os.path.join(report_dir, f"{output_filename}.pdf")

    doc.save(output_docx_path)

    try:
        convert(output_docx_path, output_pdf_path)
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

    print(f"Documento PDF generado en: {output_pdf_path}")


def traducir_texto(texto):
    """
    Traduce un texto del inglés al español utilizando Google Translator.

    Args:
        texto (str): El texto en inglés a traducir.

    Returns:
        str: El texto traducido al español.
    """
    return GoogleTranslator(source="en", target="es").translate(texto)


  
def obtener_tipos_mantenimiento(ruta_docx):
    """
    Lee un archivo DOCX y extrae los tipos de mantenimiento de la última columna de cada tabla.

    Asume que la última columna contiene una cadena con tipos de mantenimiento
    separados por comas.

    Args:
        ruta_docx (str): La ruta al archivo DOCX.

    Returns:
        list: Una lista de cadenas únicas que representan los tipos de mantenimiento encontrados.
    """
    doc = Document(ruta_docx)
    tipos_mantenimiento = set()
    
    for tabla in doc.tables:
        for fila in tabla.rows[1:]:  # Omitir la primera fila si es encabezado
            ultima_celda = fila.cells[-1].text.strip()
            tipos = [t.strip() for t in ultima_celda.split(',') if t.strip()]
            tipos_mantenimiento.update(tipos)
    
    return list(tipos_mantenimiento)
  
def ejecutar_busquedas(ruta_docx):
    """
    Ejecuta búsquedas en un portal de contratación basadas en los tipos de mantenimiento
    extraídos de un documento DOCX.

    Obtiene los tipos de mantenimiento del DOCX, los traduce al español y realiza
    búsquedas en el portal web utilizando Selenium para cada tipo traducido,
    además de una búsqueda general por 'mantenimiento'.

    Args:
        ruta_docx (str): La ruta al archivo DOCX que contiene los tipos de mantenimiento.

    Returns:
        list: Una lista de URLs encontradas durante las búsquedas.
    """
    tipos_mantenimiento = obtener_tipos_mantenimiento(ruta_docx)
    resultados_totales = []
    
    for tipo in tipos_mantenimiento:
        tipo_traducido = traducir_texto(tipo)
        print(f"Buscando: {tipo} -> {tipo_traducido}")
        resultados_totales.extend(buscar(tipo_traducido))
    
    resultados_totales.extend(buscar('mantenimiento'))

    print("resultados totales", resultados_totales)
    return resultados_totales


def buscar(mantenimiento):  
    """
    Realiza una búsqueda en el portal de contratación de la UBU utilizando Selenium.

    Navega a la página de búsqueda avanzada, llena el campo "Objeto" con el término
    de mantenimiento, marca la casilla de organismos dependientes y hace clic en buscar.
    Extrae los enlaces de los resultados encontrados.

    Args:
        mantenimiento (str): El término de mantenimiento a buscar.

    Returns:
        list: Una lista de URLs de los resultados de la búsqueda.
    """
    # El navegador no muestra su interfaz gráfica.
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Habilita el modo headless

    # Iniciar el navegador
    driver = webdriver.Chrome(options=chrome_options)
    driver.get("https://contratacion.ubu.es/licitacion/busquedaAvanzConc.do")

    # Esperar a que la página cargue
    time.sleep(3)
   
    # # Llenar el campo "Expediente"
    # campo_expediente = driver.find_element(By.ID, "expediente")  
    # campo_expediente.send_keys('UBU/2023/0018')
    try:
        # Llenar el campo "Objeto"
        campo_objeto = driver.find_element(By.ID, "ObjContrato")    
        campo_objeto.send_keys(mantenimiento)
    except Exception as e:
        print(f"Error al encontrar el campo de búsqueda: {e}")



    # Marcar la casilla "Ver expedientes de organismos dependientes"
    casilla_organismos = driver.find_element(By.NAME, "descendientes")  
    if not casilla_organismos.is_selected():
        casilla_organismos.click()  # Marcar la casilla si no está seleccionada

    # Hacer clic en el botón de búsqueda
    boton_buscar = driver.find_element(By.NAME, "busquedaFormAvanz")  
    boton_buscar.click()

     # Encuentra todas las filas de la tabla con la clase 'resultados'
    try:
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, ".resultados"))
        )
        resultados = driver.find_elements(By.CSS_SELECTOR, ".resultados tbody tr")
    except TimeoutException:
        print(f"No se encontraron resultados para {mantenimiento}")
        resultados = []

    enlaces = []  
    
    # Verifica si hay resultados
    if len(resultados) == 0:
        print("No se encontraron resultados.")
    else:
        for resultado in resultados:
            # Extraer el enlace de la primera columna (suponiendo que está en la primera columna)
            enlace = None
            enlace_elemento = resultado.find_element(By.CSS_SELECTOR, "a")
            if enlace_elemento:
                enlace = enlace_elemento.get_attribute("href")  # Obtener el atributo href del enlace
                enlaces.append(enlace)  # Guardar todos los enlaces en la 
    # Cerrar el navegador cuando termine
    driver.quit()
    return enlaces
    
    
def generar():
    """
    Función principal para generar el informe 1.19.

    Coordina el proceso completo:
    1. Obtiene la ruta del documento DOCX de mantenimiento de edificios.
    2. Ejecuta búsquedas en el portal de contratación basadas en los tipos de mantenimiento
       encontrados en el DOCX.
    3. Ejecuta la API del LLM sobre los enlaces encontrados para extraer datos estructurados.
    4. Genera el informe final en formato Word y PDF con los datos extraídos.
    """
    docx_path =os.path.join(base_dir, 'Campus_Building_Maintenance.docx')
    enlaces= ejecutar_busquedas(docx_path)
    datos = ejecutar_API(enlaces)
    generar_informe(datos)





    

