import sqlite3
import os
import subprocess
import sys
from docx import Document
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.table import Table
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from helpers import add_hyperlink, insert_table_after_paragraph, eliminar_tablas_vacias

SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# Agregar `src` al sys.path 
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)

def verificar_tipos_programa(anio, db_path):
    """
    Verifica si existen datos para los tipos de programa 'grado' y 'master' en un año dado.

    Parámetros:
    - anio (int o str): El año que se desea verificar.
    - db_path (str): Ruta a la base de datos SQLite.

    Retorna:
    - dict: Información de existencia y cantidad total para cada tipo.
            Ejemplo: {
                'grado': {'existe': True, 'cantidad_total': 100},
                'master': {'existe': False, 'cantidad_total': 0}
            }
    """
    # Verificar que la base de datos exista
    if not os.path.exists(db_path):
        raise FileNotFoundError(f"No se encontró la base de datos en {db_path}")

    # Conectar a la base de datos
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    tipos = ['grado', 'master']
    resultado = {}
    try:
        for tipo in tipos:
            cursor.execute("""
                SELECT COUNT(*)
                FROM busqueda
                WHERE anho = ? AND tipo_programa = ?
            """, (anio, tipo))
            cantidad = cursor.fetchone()[0]
            resultado[tipo] = {
                'existe': cantidad > 0,
                'cantidad': cantidad
            }

        return resultado

    finally:
        conn.close()





def fill_description(doc, anho, total):
    """
    Llena la sección de descripción del documento con texto, una tabla y texto adicional.

    Busca el párrafo que contiene "Description:" y lo reemplaza, insertando
    una tabla con el año y el total de cursos, seguida de un párrafo resumen.

    Parámetros:
    - doc (Document): El objeto Document de python-docx.
    - anho (str): El año a incluir en la descripción y tabla.
    - total (int): El número total de cursos a incluir en la descripción y tabla.
    """
    
    description_text = (
        "(Please describe the total of courses/subjects offered on your campus. "
        "The following is an example of the description. You can describe more related items if needed.)"
    )

    for para in doc.paragraphs:
        if "Description:" in para.text:
            # Reemplazar el texto con el texto de introducción
            para.clear()

            # Añadimos el texto "Description:" en negrita
            run_desc = para.add_run("Description:\n\n")
            run_desc.bold = True

            # Añadimos el texto descriptivo en cursiva
            run_italic = para.add_run(description_text)
            run_italic.italic = True

            # Insertar tabla justo después de este párrafo
            table = insert_table_after_paragraph(para, doc, 2, 2)

            # Llenar tabla
            table.cell(0, 0).text = "Year"
            table.cell(0, 1).text = "Total Courses"
            table.cell(1, 0).text = anho
            table.cell(1, 1).text = str(total)

            tbl_element = table._element
            new_para_element = OxmlElement("w:p")
            tbl_element.addnext(new_para_element)
            new_paragraph = Paragraph(new_para_element, para._parent)
            new_paragraph.add_run(f"Total number of courses offered in {anho} = {total} courses (not modules)")


            break

def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str = None) -> Paragraph:
    """
    Inserta un nuevo párrafo justo después del párrafo dado.

    Parámetros:
    - paragraph (Paragraph): El párrafo después del cual se insertará el nuevo párrafo.
    - text (str, optional): El texto inicial para el nuevo párrafo. Por defecto es "".
    - style (str, optional): El estilo para el nuevo párrafo. Por defecto es None.

    Retorna:
    - Paragraph: El objeto Paragraph recién creado.
    """
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style
    return new_paragraph

   

def generar_informe(total, anio):
    """
    Genera el informe Word y PDF para el reporte 6.2.

    Utiliza una plantilla .docx, reemplaza el encabezado, añade hipervínculos,
    llena la sección de descripción con una tabla y guarda el resultado
    en formato .docx y .pdf.

    Parámetros:
    - total (int): El número total de cursos/materias a incluir en el informe.
    - anio (str): El año para el cual se genera el informe.
    """
    
    base_dir = os.path.dirname(__file__)  # Directorio base donde se encuentra el script
    template_path = os.path.join(base_dir, 'informe_general.docx')
    
    if not os.path.exists(template_path):
        print(f"❌ Error: No se encontró la plantilla en {template_path}")
        sys.exit(1)

    output_filename = "University_Country_6_2_Total_number_of_courses_or_modules_offered"
    
    doc = Document(template_path)

    # 1. Eliminar cualquier tabla vacía en el documento
    eliminar_tablas_vacias(doc)

    # 2. Reemplazar encabezado
    for paragraph in doc.paragraphs:
        if "[6] Education and Research (ED)" in paragraph.text:
            paragraph.text = "[6] Education and Research (ED)\n\n[6.2]  Total Number of Courses/Subjects Offered\n\n"          
            # Agregar el primer enlace debajo de la sección "[6.2]"
            p1 =insert_paragraph_after(paragraph)

            enlace1 = "https://www.ubu.es/grados-ordenados-por-ramas-de-conocimiento"
            add_hyperlink(p1, enlace1, enlace1)

            # Agregar el segundo enlace debajo del primero
            p2 = insert_paragraph_after(paragraph)

            enlace2 = "https://www.ubu.es/estudios/oferta-de-estudios/masteres-universitarios-oficiales"
            add_hyperlink(p2, enlace2, enlace2)
            
            break
        

    # 4. Añadir la descripción con tabla y datos
    fill_description(doc, anio, total)

    report_dir = os.path.join(SRC_DIR, "generated_reports", "report_6_2", anio)
    os.makedirs(report_dir, exist_ok=True)  # Crea la carpeta si no existe

    output_docx_path = os.path.join(report_dir, f"{output_filename}.docx")
    output_pdf_path = os.path.join(report_dir, f"{output_filename}.pdf")

    # 5. Guardar en Word
    doc.save(output_docx_path)

    # 6. Intentar conversión a PDF
    try:
        convert(output_docx_path, output_pdf_path)
        print(f"\nDocumento PDF generado en: {output_pdf_path}")
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

 
# ------------------ FUNCIÓN PRINCIPAL ------------------
def generar(anio):
    """
    Función principal para iniciar la generación del informe 6.2.

    Verifica la existencia de datos de 'grado' y 'master' para el año especificado
    en la base de datos. Si los datos están completos, calcula el total y genera
    el informe Word y PDF. Si faltan datos, informa al usuario.

    Parámetros:
    - anio (str): El año para el cual se desea generar el informe.
    """
    
    db_path = 'instance/busqueda.db'
    resultados = verificar_tipos_programa(anio, db_path)

    if all(info['existe'] for info in resultados.values()):
        print("\nDatos completos. Generando informe...")
        # Extraer las cantidades para cada tipo
        total_grado = resultados['grado']['cantidad']
        total_master = resultados['master']['cantidad']
        total_general = total_grado + total_master  
        generar_informe(total_general, anio)
            
    else:
        print("\nFaltan datos para alguno de los tipos de programa:\n")
        for tipo, info in resultados.items():
            if not info['existe']:
                print(f"- No hay datos para '{tipo}'. Descargue los datos faltantes.")

                    
       

  


    

