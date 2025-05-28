
import json
import os
import sys
import re
from nbformat import convert
import requests
from docx import Document
from dotenv import load_dotenv
from config import cargar_configuracion


SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Agregar `src` al sys.path para poder importar `config.py`
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)


# Variable global
base_dir = os.path.dirname(__file__)
load_dotenv()
# Configuración de la API
config = cargar_configuracion()
base_url = config["base_url"]
api_key = config["api_key"]
myModel = config["model"]


carpeta = os.path.join(base_dir, 'UBU_Verde_informes')




def extraer_anhos(year_range):
    """
    Extrae los años académicos desde un string tipo 'YYYY-YYYY'.

    Convierte un rango de años como "2022-2025" en una lista de strings
    representando años académicos, como ["2022-23", "2023-24", "2024-25"].

    Args:
        year_range (str): El rango de años en formato "YYYY-YYYY".

    Returns:
        list[str]: Una lista de strings representando los años académicos.
                   Retorna una lista vacía si el formato de entrada no coincide.
    """
    match = re.match(r"(\d{4})-(\d{4})", year_range)
    if match:
        inicio, fin = int(match.group(1)), int(match.group(2))
        
        return [f"{año}-{str(año+1)[-2:]}" for año in range(inicio, fin)]
    return []




def buscar_ficheros(cursos, carpeta):
    """
    Busca archivos .docx dentro de una carpeta que contengan nombres de cursos.

    Busca archivos que terminen con .docx y cuyo nombre contenga el formato
    del curso (ej. "2022-23" o "202223").

    Args:
        cursos (list[str]): Lista de strings representando los años académicos (ej. ["2022-23", "2023-24"]).
        carpeta (str): La ruta a la carpeta donde buscar los archivos.

    Returns:
        list[tuple[str, str]]: Una lista de tuplas, donde cada tupla contiene
                                el nombre del curso encontrado y la ruta completa al archivo.
    """
    archivos_encontrados = []
    for archivo in os.listdir(carpeta):
        for curso in cursos:
            curso_formato = curso.replace("-", "")  # Eliminar el guion para coincidir con 22-23 y 2022-23
            if re.search(rf"{curso}|\b{curso_formato}\b", archivo) and archivo.endswith(".docx"):
                archivos_encontrados.append((curso, os.path.join(carpeta, archivo)))
    return archivos_encontrados



def leer_contenido_docx(ruta_fichero):
    """
    Lee el texto completo de un archivo Word (.docx).

    Args:
        ruta_fichero (str): La ruta completa al archivo .docx.

    Returns:
        str | None: El texto extraído del documento, o None si ocurre un error.
    """
    try:
        doc = Document(ruta_fichero)
        texto = "\n".join([p.text for p in doc.paragraphs])
        return texto
    except Exception as e:
        print(f"Error al leer {ruta_fichero}: {e}")
        return None


def fill_description(doc, year_range, resultados):
    """
    Llena la sección de descripción en el documento Word con los resultados del conteo.

    Calcula el total y el promedio de actividades sostenibles por año y construye
    un texto descriptivo que incluye estos datos y el desglose por año académico.
    Reemplaza el párrafo que contiene "Description:" con este texto.

    Args:
        doc (Document): Objeto Document de python-docx.
        year_range (str): El rango de años utilizado para la búsqueda (ej. "2022-2025").
        resultados (dict[str, int]): Un diccionario donde las claves son los años
                                     académicos y los valores son el número de
                                     actividades sostenibles encontradas para ese año.
    """
    
    # Calcular el número total de años y el total de actividades sostenibles
    num_years = len(resultados)
    if resultados:
        total = sum(resultados.values())
    else:
        total = 0
    
    if num_years > 0:
        average = total // num_years 
    else:
       average = 0

    # Construcción de la descripción
    description_text = (
        f"Events related to environment and sustainability hosted or organized by the University of Burgos "
        f"in the academic year {year_range}\n\n"
        f"Total number of sustainability/environment related events in:\n\n"
    )

    # Agregar cada año y su número de actividades
    for year, count in resultados.items():
        description_text += f"{year}: {count}.\n\n"

    description_text += (
        f"A total average per annum over the last {num_years} years is {average} events "
        f"(e.g. conferences, workshops, awareness raising, practical training, etc.).\n\n"
    )

    # Reemplazar el texto en el documento
    for para in doc.paragraphs:
        if "Description:" in para.text:
            para.text = f"Description:\n\n{description_text}"
            break




 
def extraer_datos_llm(texto):
    """
    Extrae el número de actividades relacionadas con la sostenibilidad de un texto usando un modelo LLM.

    Envía el texto a un modelo de lenguaje a través de una API y le pide que cuente
    las actividades relacionadas con los ODS.

    Args:
        texto (str): El texto que contiene la lista de actividades.

    Returns:
        int | None: El número de actividades sostenibles encontradas, o None si
                    no se pudo extraer un número válido de la respuesta del modelo.
    """
    if not texto:
       return None

    body = {
        "model": myModel,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You will receive a text containing a list of activities. "
                    "Identify which activities are related to environmental sustainability, based on if they align with: "
                    "1. Fin de la pobreza\n"
                    "2. Hambre cero\n"
                    "3. Salud y bienestar\n"
                    "4. Educación de calidad\n"
                    "5. Igualdad de género\n"
                    "6. Agua limpia y saneamiento\n"
                    "7. Energía asequible y no contaminante\n"
                    "8. Trabajo decente y crecimiento económico\n"
                    "9. Industria, innovación e infraestructura\n"
                    "10. Reducción de las desigualdades\n"
                    "11. Ciudades y comunidades sostenibles\n"
                    "12. Producción y consumo responsables\n"
                    "13. Acción por el clima\n"
                    "14. Vida submarina\n"
                    "15. Vida de ecosistemas terrestres\n"
                    "16. Paz, justicia e instituciones sólidas\n"
                    "17. Alianzas para lograr los objetivos\n"
                    "Count only those activities and return the number.\n\n"
                    "Return ONLY the number, with no extra text, explanation, or formatting."
                )
            },
            {
                "role": "user",
                "content": texto
            }
        ],
        "temperature": 0.2
    }
        
    body_json = json.dumps(body)

    response = requests.post(
        f"{base_url}/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        data=body_json
    )    

    #  Obtener respuesta limpia
    message_content = response.json().get('choices', [{}])[0].get('message', {}).get('content', '').strip()
        
    #  Eliminar `</think>` si existe
    if "</think>" in message_content:
        message_content = message_content.split("</think>")[-1].strip()
    match = re.search(r"\d+", message_content)
    if match:
        return int(match.group())
    print(f"No se encontró un número válido en la respuesta del modelo:\n{message_content}")
    return None

       

def procesar_documentos(rango_cursos, carpeta):
    """
    Procesa los documentos de Word para un rango de años académicos y extrae el conteo de actividades sostenibles.

    Busca los archivos de Word correspondientes a los años académicos especificados,
    lee el contenido de cada archivo y utiliza `extraer_datos_llm` para contar
    las actividades sostenibles en cada uno.

    Args:
        rango_cursos (str): El rango de años en formato "YYYY-YYYY".
        carpeta (str): La ruta a la carpeta que contiene los archivos de Word.

    Returns:
        dict[str, int]: Un diccionario donde las claves son los años académicos
                        y los valores son el número de actividades sostenibles
                        encontradas para ese año.
    """
    cursos = extraer_anhos(rango_cursos)
    archivos = buscar_ficheros(cursos, carpeta)
    
    resultados = {}
    for curso, ruta in archivos:
        texto = leer_contenido_docx(ruta)
        num_actividades = extraer_datos_llm(texto)
        resultados[curso] = num_actividades

    return resultados
  
     

# Función que genera el informe
def generar(anho):
    """
    Genera el informe Word y PDF para el reporte 6.8.

    Procesa los documentos de Word para el rango de años especificado, obtiene
    los resultados del conteo de actividades sostenibles, carga una plantilla
    de documento Word, reemplaza el encabezado y la descripción con los datos
    obtenidos, guarda el documento modificado como .docx, y luego intenta
    convertirlo a .pdf.

    Args:
        anho (str): El rango de años en formato "YYYY-YYYY" para el cual se
                    generará el informe.
    """
    
    resultados = procesar_documentos(anho, carpeta)


    template_path = os.path.join(base_dir, 'informe_general.docx')
   
    if not os.path.exists(template_path):
        print(f"Error: No se encontró la plantilla {template_path}")
        return

    output_filename = "University_Country_6_8_Number_of_events_related_to_environment_and_sustainability"
    
    doc = Document(template_path)

    # Reemplazar texto en la plantilla
    for paragraph in doc.paragraphs:
        if  "[6] Education and Research (ED)" in paragraph.text:
            paragraph.text = f" [6] Education and Research (ED) \n\n [6.8] Number of Events Related to Sustainability"   
            break
   

    fill_description(doc,anho, resultados)

    report_dir = os.path.join(SRC_DIR, "generated_reports", "report_6_8", anho)
    os.makedirs(report_dir, exist_ok=True)  # Crea la carpeta si no existe

    output_docx_path = os.path.join(report_dir, f"{output_filename}.docx")
    output_pdf_path = os.path.join(report_dir, f"{output_filename}.pdf")

    doc.save(output_docx_path)

    try:
        convert(output_docx_path, output_pdf_path)
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

    print(f"Documento PDF generado en: {output_pdf_path}")

 


    

