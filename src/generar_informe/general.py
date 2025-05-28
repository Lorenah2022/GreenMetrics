import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import sys
import importlib  # Importar dinámicamente el código de los informes
import tkinter as tk
from tkinter import ttk

def crear_word_documento():
    """
    Crea el documento Word utilizando la plantilla 'modelo_base.docx',
    rellena los datos generales de la universidad y reemplaza el marcador
    de posición del logo con una imagen.

    La plantilla 'modelo_base.docx' y la imagen del logo 'logo-UBU.jpg'
    deben estar en ubicaciones relativas esperadas.

    Sale del script si la plantilla no se encuentra.
    """
    base_dir = os.path.dirname(__file__)  # Directorio base donde se encuentra el script
    template_path = os.path.join(base_dir, 'modelo_base.docx')
    output_docx_path = os.path.join(base_dir, f"informe_general.docx")
    logo_path = os.path.join(base_dir, '../static/images/logo-UBU.jpg')

    if not os.path.exists(template_path):
        print(f"Error: No se encontró la plantilla {template_path}")
        sys.exit(1)

    university_name = "University of Burgos"
    country_name = "Spain"
    web_address = "https://www.ubu.es/"

    def cm_to_inches(cm):
        """
        Convierte una medida de centímetros a pulgadas.

        Args:
            cm (float): Medida en centímetros.

        Returns:
            float: Medida convertida en pulgadas.
        """
        return cm / 2.54
    
    # Función para insertar el logo en lugar del texto "Put your university logo here"
    def replace_logo_placeholder(doc, logo_path):
        """
        Busca el texto "Put your university logo here" en la cabecera del documento
        y lo reemplaza con la imagen del logo especificada.

        Args:
            doc (docx.document.Document): El objeto documento de python-docx.
            logo_path (str): Ruta al archivo de imagen del logo.
        """
        header = doc.sections[0].header  # Acceder a la cabecera de la primera sección

        for para in header.paragraphs:
            if "Put your university logo here" in para.text:
                # Encuentra y elimina el texto del marcador de posición
                inline_shapes = para._element.xpath(".//w:t[text()='Put your university logo here']")
                if inline_shapes:
                    inline_shapes[0].getparent().remove(inline_shapes[0])  # Elimina solo el texto del marcador
                    # Añade la imagen en su lugar con dimensiones específicas
                    run = para.add_run()
                    run.add_picture(logo_path, width=Inches(cm_to_inches(3.06)), height=Inches(cm_to_inches(2.25)))
                break
    
    def fill_university_info(doc):
        '''
        Rellena los marcadores de posición de información de la universidad en un objeto documento.

        Itera a través de los párrafos del objeto documento proporcionado y reemplaza
        textos marcadores de posición específicos ("University", "Country", "Web Address") con
        valores correspondientes de variables globales o accesibles (university_name, country_name, web_address).

        Args:
            doc (docx.document.Document): El objeto documento (e.g., python-docx Document) a modificar.
        '''
        for para in doc.paragraphs:
            if "University" in para.text:
                para.text = f"University: {university_name}"
            elif "Country" in para.text:
                para.text = f"Country: {country_name}"
            elif "Web Address" in para.text:
                para.text = f"Web Address: {web_address}"

    # Crear el documento con la plantilla
    doc = Document(template_path)
    replace_logo_placeholder(doc, logo_path)
    fill_university_info(doc)
    doc.save(output_docx_path)
    print(f"Documento Word creado en: {output_docx_path}")

def ejecutar_informe_especifico(anho, informe, excel):
    """
    Ejecuta un informe específico importando dinámicamente el módulo correspondiente.

    Basado en el identificador del informe proporcionado, importa el módulo Python
    asociado y llama a su función 'generar', pasando los argumentos necesarios
    (año o datos de Excel).

    Args:
        anho (int or str): El año relevante para la generación del informe.
                           Utilizado por los informes 6_1, 6_2, 6_3, 6_7, 6_8.
        informe (str): El identificador del informe a ejecutar (ej. "1_19", "6_1").
        excel (any): Datos o ruta relacionados con un archivo Excel.
                     Utilizado por el informe 6_4.

    Raises:
        ImportError: Si el módulo del informe especificado no puede ser importado.
        Exception: Captura cualquier otra excepción durante la importación
                   e imprime un mensaje de error.
    """
    try:
         # Importa dinámicamente los ficheros py
        if  informe == "1_19":
            informe_1_19 = importlib.import_module('informe_1_19') 
            informe_1_19.generar()
        elif informe == "6_1":
            informe_6_1 = importlib.import_module('informe_6_1') 
            informe_6_1.generar(anho)  
        elif informe == "6_2":
            informe_6_2 = importlib.import_module('informe_6_2') 
            informe_6_2.generar(anho) 
        elif informe == "6_3":
            informe_6_3 = importlib.import_module('informe_6_3') 
            informe_6_3.generar(anho) 
        elif informe == "6_4":
            informe_6_4 = importlib.import_module('informe_6_4') 
            informe_6_4.generar(excel) 
        elif informe == "6_7":
            informe_6_7 = importlib.import_module('informe_6_7')  
            informe_6_7.generar(anho)
        elif informe == "6_8":
            informe_6_8 = importlib.import_module('informe_6_8') 
            informe_6_8.generar(anho)
       
    except Exception as e:
        print(f"Error al cargar informe_{informe}.py: {e}")

def generar_informe(anho, informe,excel):
    """
    Función principal que coordina la creación del documento Word general
    y la ejecución de un informe específico.

    Primero crea el documento Word base y luego llama a la función
    para ejecutar el informe particular solicitado.

    Args:
        anho (int or str): El año relevante para la generación del informe.
        informe (str): El identificador del informe a ejecutar.
        excel (any): Datos o ruta relacionados con un archivo Excel, si es necesario para el informe.
    """
    crear_word_documento()
    ejecutar_informe_especifico(anho, informe,excel)

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Uso correcto: python script.py <anho> <informe><excel>")
        sys.exit(1)
    anho = sys.argv[1]
    informe = sys.argv[2]
    excel = sys.argv[3]
    generar_informe(anho, informe, excel)
