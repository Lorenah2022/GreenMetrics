
import os
import sys
from io import BytesIO
import re
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.chrome.options import Options


# Variable global
base_dir = os.path.dirname(__file__)

SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# Agregar `src` al sys.path 
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)

def generar_enlace (year_range) :
    """
    Genera el enlace de búsqueda de Google Scholar para un rango de años y palabras clave.

    Args:
        year_range (str): Rango de años en formato "YYYY-YYYY".

    Returns:
        tuple: Una tupla que contiene:
            - str: El enlace de búsqueda generado.
            - int: El año de inicio del rango.
            - int: El año de fin del rango.
    """
    start_year, end_year = map(int, year_range.split('-'))
    enlace = f"https://scholar.google.es/scholar?hl=es&as_sdt=0%2C5&as_ylo={start_year}&as_yhi={end_year}&q=%22Universidad+de+Burgos%22+%26+%28%22green%22+OR+%22environment%22+OR+%22sustainability%22+OR+%22renewable+energy%22+OR+%22climate+change%22%29&btnG="
    return enlace, start_year, end_year
            
def fill_description(doc, year_range, num_courses):
    """
    Llena la sección de descripción en el documento Word con información sobre la búsqueda.

    Busca los párrafos que contienen "Description:" y "Additional evidence link",
    reemplaza el texto de la descripción con los detalles de la búsqueda y el número
    de resultados, y añade un hipervínculo al enlace de búsqueda.

    Args:
        doc (Document): Objeto Document de python-docx.
        year_range (str): Rango de años utilizado en la búsqueda (ej. "2020-2022").
        num_courses (int): Número total de resultados encontrados.
    """
    enlace, start_year, end_year=generar_enlace(year_range)
    num_years = end_year - start_year + 1  # Calcular el número de años en el rango

    description_text = (
        f"Results obtained from the following query in the {num_years}-year interval set and using the keywords provided in the questionnaire.\n\n"
        f"\"Universidad de Burgos\" & (\"green\" OR \"environment\" OR \"sustainability\" OR \"renewable energy\" OR \"climate change\")\n\n"
        f"Scholarly publications on sustainability in the academic years {start_year}-{end_year}.\n\n"
        f"A total average per annum over the last {num_years} years of {num_courses} / {num_years} = {num_courses // num_years} publications.\n\n"
        
    )
    
    Additional_evidence_text=(
        f"Additional evidence link (i.e., for videos, more images, or other files that are not included in this file):"
    )

    for para in doc.paragraphs:
        if "Description:" in para.text:
            para.text = f"Description:\n\n{description_text}"
            
        elif f"{Additional_evidence_text}" in para.text:
            add_hyperlink(para, enlace, "Link")
            break

def agregar_imagen_a_tabla(doc, captura_pantalla):
    """
    Agrega una imagen (captura de pantalla) a la primera tabla del documento Word.

    Inserta la imagen en la primera celda de la primera fila de la primera tabla
    encontrada en el documento y ajusta su tamaño. Si la tabla tiene una segunda
    fila, añade un texto descriptivo en la primera celda de esa fila; de lo contrario,
    añade una nueva fila con el texto.

    Args:
        doc (Document): Objeto Document de python-docx.
        captura_pantalla (bytes): Datos binarios de la imagen de la captura de pantalla.
    """
    
    # Buscar la tabla en el documento (se asume que es la primera tabla en el documento)
    table = doc.tables[0]  # Suponiendo que es la primera tabla en el documento
    rows = table.rows
    
    # Insertar la captura de pantalla en la primera celda de la primera fila
    cell = rows[0].cells[0]
    image_stream = BytesIO(captura_pantalla)  # Convertir los bytes a un stream
    picture = cell.paragraphs[0].add_run().add_picture(image_stream)  # Añadir la imagen
    
    # Establecer un tamaño específico para la imagen si es necesario
    # Aquí ajustas el tamaño de la imagen según lo necesites (ejemplo: 2 pulgadas de ancho)
    picture.width = Inches(6)
    picture.height = Inches(3)  # Si también necesitas ajustar la altura

    # Después de insertar la imagen, la celda se ajustará automáticamente en tamaño para contenerla

    # Si la tabla ya tiene dos filas, vamos a insertar el texto en la segunda celda de la segunda fila
    if len(rows) > 1:
        text_cell = rows[1].cells[0]  # Asumimos que la segunda fila tiene la primera celda para el texto
        text_cell.text = 'Scholarly publications on sustainability (Universidad de Burgos, Spain)'

    else:
        # Si no hay suficientes filas, añadimos una nueva fila con dos celdas y el texto
        new_row = table.add_row()
        new_row.cells[0].text = 'Scholarly publications on sustainability (Universidad de Burgos, Spain)'
        # La segunda celda estará vacía (podrías agregar algo más si lo necesitas)
        new_row.cells[1].text = ''


def add_hyperlink(paragraph, url, text):
    """
    Agrega un hipervínculo a un párrafo en un documento Word.

    Configura el estilo del hipervínculo (color azul y subrayado) y añade el texto
    visible del enlace al párrafo.

    Args:
        paragraph (docx.text.paragraph.Paragraph): El objeto párrafo de python-docx.
        url (str): La URL de destino del hipervínculo.
        text (str): El texto visible del hipervínculo.
    """
    # Crear la relación del hipervínculo en el documento
    part = paragraph._element
    rId = paragraph._parent.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rId)

    # Crear el run del enlace
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Color azul
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # Azul
    rPr.append(color)

    # Subrayado manual
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")  # Subrayado
    rPr.append(u)

    new_run.append(rPr)

    # Agregar el texto visible
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    part.append(hyperlink)

def generar_informe(captura_pantalla,anho,numero_resultados):
    """
    Genera el informe Word y PDF para el reporte 6.7.

    Carga una plantilla de documento Word, reemplaza el encabezado específico
    del reporte 6.7, agrega la captura de pantalla a la tabla, llena la sección
    de descripción, guarda el documento modificado como .docx, y luego intenta
    convertirlo a .pdf.

    Args:
        captura_pantalla (bytes): Datos binarios de la imagen de la captura de pantalla.
        anho (str): El rango de años para el cual se generó el informe (ej. "2020-2022").
        numero_resultados (int): El número total de resultados encontrados en la búsqueda.
    """    
        
    template_path = os.path.join(base_dir, 'informe_general.docx')
   
    if not os.path.exists(template_path):
        print(f"Error: No se encontró la plantilla {template_path}")
        return

    output_filename = "University_Country_6_7_Number_of_scholarly_publications_on_sustainability_UBU"
    

    doc = Document(template_path)

    # Reemplazar texto en la plantilla
    for paragraph in doc.paragraphs:
        if  "[6] Education and Research (ED)" in paragraph.text:
            paragraph.text = f" [6] Education and Research (ED) \n\n [6.7] Number of scholarly publications on sustainability"   
            break
   

    agregar_imagen_a_tabla(doc, captura_pantalla)

    fill_description(doc,anho,numero_resultados)

    report_dir = os.path.join(SRC_DIR, "generated_reports", "report_6_7", anho)
    os.makedirs(report_dir, exist_ok=True)  # Crea la carpeta si no existe

    output_docx_path = os.path.join(report_dir, f"{output_filename}.docx")
    output_pdf_path = os.path.join(report_dir, f"{output_filename}.pdf")

    doc.save(output_docx_path)

    try:
        convert(output_docx_path, output_pdf_path)
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

    print(f"Documento PDF generado en: {output_pdf_path}")


def buscar(anho):
    """
    Realiza una búsqueda en Google Scholar para un rango de años y palabras clave.

    Utiliza Selenium para abrir un navegador (en modo headless), navega al enlace
    de búsqueda generado, espera a que carguen los resultados, extrae el número
    total de resultados y toma una captura de pantalla.

    Args:
        anho (str): El rango de años para la búsqueda (ej. "2020-2022").

    Returns:
        tuple: Una tupla que contiene:
            - bytes: Datos binarios de la captura de pantalla.
            - int: El número total de resultados encontrados.
    """
    
    # Configurar Selenium en modo headless
    chrome_options = Options()
    chrome_options.add_argument("--headless")  

    # Iniciar el navegador
    driver = webdriver.Chrome(options=chrome_options)
    enlace, start_year, end_year = generar_enlace(anho)
    driver.get(enlace)

    try:
        # Esperar hasta que el elemento que contiene el número de resultados esté presente
        resultado_elemento = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "div#gs_ab_md div.gs_ab_mdw"))
        )
        
        # Extraer el texto del elemento que contiene el número de resultados
        texto = resultado_elemento.text
        # Buscar el número de resultados con un patrón
        numero_resultados = re.search(r"Aproximadamente (\d{1,3}(?:\.\d{3})*) resultados", texto)

        if numero_resultados:
            numero_resultados = numero_resultados.group(1).replace(".", "")  # Eliminar los puntos y obtener el número limpio
        else:
            print("No se encontró el número de resultados.")
        

        # Captura de pantalla
        captura_pantalla = driver.get_screenshot_as_png()
        
    except Exception as e:
        print("Error al extraer los resultados:", e)
    
    print("resultados",numero_resultados)

    numero_resultados_convertido=int(numero_resultados)
    # Cerrar el navegador
    driver.quit()
    return captura_pantalla, numero_resultados_convertido
    
 


def generar(anho):
    """
    Función principal para generar el reporte 6.7.

    Realiza la búsqueda en Google Scholar para obtener la captura de pantalla
    y el número de resultados, y luego llama a la función para generar el informe
    Word y PDF.

    Args:
        anho (str): El rango de años para el cual se generará el informe (ej. "2020-2022").
    """
    captura_pantalla, numero_resultados= buscar(anho)
    generar_informe(captura_pantalla,anho,numero_resultados)




    

